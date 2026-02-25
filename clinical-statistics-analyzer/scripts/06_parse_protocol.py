#!/usr/bin/env python3
"""
06_parse_protocol.py
Protocol Document Parser for Clinical Trial Protocols

Parses clinical trial protocol documents (DOCX/PDF) and extracts:
- Study metadata (ID, title, phase, sponsor)
- Disease/indication information
- Study design (randomization, blinding, arms)
- Endpoints (primary, secondary)
- Treatment arms
- Inclusion/exclusion criteria
- Sample size and statistical considerations

Author: Clinical Statistics Analyzer
Version: 1.0.0
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any

# Try to import required libraries
try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import pdfplumber
except ImportError:
    print("Installing pdfplumber...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"])
    import pdfplumber


class ProtocolParser:
    """
    Parser for clinical trial protocol documents.
    Supports DOCX and PDF formats.
    """
    
    def __init__(self, file_path: str):
        """
        Initialize the parser with a protocol file.
        
        Args:
            file_path: Path to the protocol document (DOCX or PDF)
        """
        self.file_path = Path(file_path)
        self.file_ext = self.file_path.suffix.lower()
        self.document_text = ""
        self.parsed_data = {}
        
    def parse(self) -> Dict[str, Any]:
        """
        Parse the protocol document and extract structured data.
        
        Returns:
            Dictionary containing parsed protocol information
        """
        if not self.file_path.exists():
            raise FileNotFoundError(f"Protocol file not found: {self.file_path}")
        
        # Extract text based on file type
        if self.file_ext == '.docx':
            self.document_text = self._parse_docx()
        elif self.file_ext == '.pdf':
            self.document_text = self._parse_pdf()
        else:
            raise ValueError(f"Unsupported file format: {self.file_ext}")
        
        # Extract structured information
        self.parsed_data = {
            "metadata": self._extract_metadata(),
            "study_design": self._extract_study_design(),
            "disease_info": self._extract_disease_info(),
            "endpoints": self._extract_endpoints(),
            "treatment_arms": self._extract_treatment_arms(),
            "eligibility": self._extract_eligibility(),
            "statistics": self._extract_statistics(),
            "raw_text": self.document_text[:10000]  # First 10k chars for reference
        }
        
        return self.parsed_data
    
    def _parse_docx(self) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(self.file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        return "\n".join(text_parts)
    
    def _parse_pdf(self) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n".join(text_parts)
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract study metadata."""
        text = self.document_text
        
        # Study ID pattern
        study_id_pattern = r'(?:Protocol|Sponsor|Invesigational)\s*(?:ID|Number|No\.?)\s*:?\s*([A-Z0-9\-]+)'
        study_id = self._search_pattern(study_id_pattern, text, group=1)
        
        # Title pattern
        title_pattern = r'(?:Title|Study\s*Title)[:\s]+(.+?)(?:\n|$)'
        title = self._search_pattern(title_pattern, text, group=1)
        
        # Phase pattern
        phase_pattern = r'(?:Phase)\s*(?:Study)?[:\s]*(I{1,3}|IV|1|2|3|4|a|b)'
        phase = self._search_pattern(phase_pattern, text, group=1)
        
        # Sponsor
        sponsor_pattern = r'(?:Sponsor)[:\s]+(.+?)(?:\n|$)'
        sponsor = self._search_pattern(sponsor_pattern, text, group=1)
        
        # Date patterns
        date_pattern = r'(?:Date|Ver(?:sion)?\.?\s*\d*)[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})'
        date = self._search_pattern(date_pattern, text, group=1)
        
        return {
            "study_id": study_id or "Unknown",
            "title": title or "Unknown",
            "phase": phase or "Unknown",
            "sponsor": sponsor or "Unknown",
            "version_date": date or "Unknown",
            "parsed_date": datetime.now().isoformat()
        }
    
    def _extract_study_design(self) -> Dict[str, Any]:
        """Extract study design information."""
        text = self.document_text
        
        # Study type
        study_type = self._search_pattern(
            r'(?:Study\s*Design|Type\s*of\s*Study)[:\s]+(.+?)(?:\n|$)',
            text, group=1
        )
        
        # Randomization
        randomization = self._search_pattern(
            r'(?:Randomization|Randomised)[:\s]*(.+?)(?:\n|$)',
            text, group=1
        )
        
        # Blinding
        blinding = self._search_pattern(
            r'(?:Blinding|Blinded|Double[\-\s]?blind|Single[\-\s]?blind)[:\s]*(.+?)(?:\n|$)',
            text, group=1
        )
        
        # Duration
        duration = self._search_pattern(
            r'(?:Duration|Treatment\s*Period|Study\s*Duration)[:\s]*(\d+\s*(?:weeks?|months?|years?|cycles?)?)',
            text, group=1
        )
        
        return {
            "study_type": study_type or "Unknown",
            "randomization": randomization or "Not specified",
            "blinding": blinding or "Not specified",
            "duration": duration or "Not specified"
        }
    
    def _extract_disease_info(self) -> Dict[str, Any]:
        """Extract disease/indication information."""
        text = self.document_text
        
        # Disease/Indication
        disease = self._search_pattern(
            r'(?:Disease|Indication|Patient\s*Population)[:\s]+(.+?)(?:\n|$)',
            text, group=1
        )
        
        # If not found in specific pattern, search for common disease patterns
        if not disease:
            disease = self._search_pattern(
                r'(?:Acute\s*Myeloid\s*Leukemia|AML|Acute\s*Lymphoblastic\s*Leukemia|ALL)',
                text
            )
        
        return {
            "disease": disease or "Unknown",
            "disease_category": self._categorize_disease(disease)
        }
    
    def _categorize_disease(self, disease: str) -> str:
        """Categorize the disease type."""
        if not disease:
            return "Unknown"
        
        disease_lower = disease.lower()
        
        if 'aml' in disease_lower or 'myeloid' in disease_lower:
            return "AML"
        elif 'all' in disease_lower or 'lymphoblastic' in disease_lower:
            return "ALL"
        elif 'cml' in disease_lower or 'myelogenous' in disease_lower:
            return "CML"
        elif 'lymphoma' in disease_lower:
            return "Lymphoma"
        elif 'myeloma' in disease_lower:
            return "Myeloma"
        else:
            return "Other"
    
    def _extract_endpoints(self) -> Dict[str, Any]:
        """Extract endpoint information."""
        text = self.document_text
        
        # Primary endpoint
        primary_patterns = [
            r'(?:Primary\s*Endpoint|Primary\s*Objective)[:\s]+(.+?)(?:\n\n|\n[A-Z])',
            r'(?:Primary)[:\s]+(.+?)(?:\n\n|\n[A-Z])'
        ]
        primary = self._search_patterns(primary_patterns, text, group=1)
        
        # Secondary endpoints
        secondary_patterns = [
            r'(?:Secondary\s*Endpoints|Secondary\s*Objectives)[:\s]+(.+?)(?:\n\n|\n[A-Z])',
            r'(?:Secondary)[:\s]+(.+?)(?:\n\n|\n[A-Z])'
        ]
        secondary = self._search_patterns(secondary_patterns, text, group=1)
        
        return {
            "primary": primary or "Unknown",
            "secondary": secondary or "Not specified"
        }
    
    def _extract_treatment_arms(self) -> List[Dict[str, Any]]:
        """Extract treatment arm information."""
        text = self.document_text
        arms = []
        
        # Pattern for treatment arms
        arm_pattern = r'(?:Arm|Treatment\s*Arm|Group)\s*[:\s]*(?:(\d+)|([A-Z])|(.+?))(?:\n|;)'
        
        # Look for common treatment arm structures
        arm_sections = re.finditer(
            r'(?:Treatment\s*Arms?|Study\s*Arms?|Cohorts?)[:\s]+(.+?)(?:\n\n|\n(?:Inclusion|Exclusion|Primary|Secondary|Statistical|$))',
            text,
            re.DOTALL | re.IGNORECASE
        )
        
        for match in arm_sections:
            arm_text = match.group(1)
            
            # Extract individual arms
            individual_arms = re.split(r'(?:\n(?:\d+\.?)|(?<!\d)\n[A-Z]\.?|\n•|\n-)', arm_text)
            
            for arm in individual_arms:
                arm = arm.strip()
                if arm and len(arm) > 5:
                    arms.append({
                        "name": arm[:100],
                        "description": arm,
                        "type": self._classify_arm_type(arm)
                    })
        
        return arms if arms else [{"name": "Not specified", "description": "Not specified", "type": "Unknown"}]
    
    def _classify_arm_type(self, arm_text: str) -> str:
        """Classify the treatment arm type."""
        text_lower = arm_text.lower()
        
        if 'placebo' in text_lower:
            return "Placebo"
        elif 'control' in text_lower:
            return "Control"
        elif 'standard' in text_lower or 'soc' in text_lower:
            return "Standard of Care"
        elif 'low' in text_lower or 'dose' in text_lower:
            return "Low Dose"
        elif 'high' in text_lower or 'dose' in text_lower:
            return "High Dose"
        elif 'combination' in text_lower:
            return "Combination"
        elif 'mono' in text_lower or 'single' in text_lower:
            return "Monotherapy"
        else:
            return "Experimental"
    
    def _extract_eligibility(self) -> Dict[str, Any]:
        """Extract eligibility criteria."""
        text = self.document_text
        
        # Inclusion criteria
        inclusion_pattern = r'(?:Inclusion\s*Criteria|Inclusion\s*Eligibility)[:\s]+(.+?)(?:\n(?:Exclusion|$))'
        inclusion = self._search_pattern(inclusion_pattern, text, group=1)
        
        # Exclusion criteria
        exclusion_pattern = r'(?:Exclusion\s*Criteria|Exclusion\s*Eligibility)[:\s]+(.+?)(?:\n(?:Study|$))'
        exclusion = self._search_pattern(exclusion_pattern, text, group=1)
        
        return {
            "inclusion": inclusion or "Not specified",
            "exclusion": exclusion or "Not specified"
        }
    
    def _extract_statistics(self) -> Dict[str, Any]:
        """Extract statistical information."""
        text = self.document_text
        
        # Sample size
        sample_size = self._search_pattern(
            r'(?:Sample\s*Size|Number\s*of\s*Patients|Total\s*Patients)[:\s]*(\d+)',
            text, group=1
        )
        
        # Power
        power = self._search_pattern(
            r'(?:Power|Statistical\s*Power)[:\s]*(\d+\%|\d+\.?\d*)',
            text, group=1
        )
        
        # Alpha/Significance
        alpha = self._search_pattern(
            r'(?:Alpha|Significance\s*Level)[:\s]*(0\.\d+)',
            text, group=1
        )
        
        return {
            "sample_size": int(sample_size) if sample_size else None,
            "power": power or "Not specified",
            "alpha": alpha or "Not specified"
        }
    
    def _search_pattern(self, pattern: str, text: str, group: int = 0) -> Optional[str]:
        """Search for a pattern in text."""
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            if group > 0:
                return match.group(group).strip()
            return match.group(0).strip()
        return None
    
    def _search_patterns(self, patterns: List[str], text: str, group: int = 0) -> Optional[str]:
        """Search for multiple patterns in text, return first match."""
        for pattern in patterns:
            result = self._search_pattern(pattern, text, group)
            if result:
                return result
        return None
    
    def save_json(self, output_path: Optional[str] = None) -> str:
        """
        Save parsed data to JSON file.
        
        Args:
            output_path: Path for output JSON file. If None, uses default naming.
            
        Returns:
            Path to saved JSON file
        """
        if not self.parsed_data:
            raise ValueError("No parsed data available. Run parse() first.")
        
        if output_path is None:
            output_path = self.file_path.parent / f"{self.file_path.stem}_parsed.json"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.parsed_data, f, indent=2, ensure_ascii=False)
        
        return str(output_path)


def main():
    """Main function for command-line usage."""
    if len(sys.argv) < 2:
        print("Usage: python 06_parse_protocol.py <protocol_file.docx/pdf> [output_json]")
        print("\nExample:")
        print("  python 06_parse_protocol.py protocol.docx")
        print("  python 06_parse_protocol.py protocol.pdf output.json")
        sys.exit(1)
    
    protocol_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        print(f"Parsing protocol: {protocol_file}")
        parser = ProtocolParser(protocol_file)
        data = parser.parse()
        
        # Print summary
        print("\n" + "="*60)
        print("Protocol Parsed Successfully")
        print("="*60)
        print(f"Study ID: {data['metadata']['study_id']}")
        print(f"Title: {data['metadata']['title']}")
        print(f"Phase: {data['metadata']['phase']}")
        print(f"Disease: {data['disease_info']['disease']}")
        print(f"Primary Endpoint: {data['endpoints']['primary']}")
        print(f"Treatment Arms: {len(data['treatment_arms'])}")
        print(f"Sample Size: {data['statistics']['sample_size']}")
        print("="*60)
        
        # Save to file
        output_path = parser.save_json(output_file)
        print(f"\nSaved to: {output_path}")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
