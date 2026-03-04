#!/usr/bin/env python3
"""
crf_spec_parser.py
CRF (Case Report Form) Specification Parser

Parses CRF specification documents (DOCX/XLSX) and extracts:
- Variable names and labels
- Data types and formats
- Valid ranges and value lists
- Skip patterns and dependencies
- Section/category information

Author: Clinical Statistics Analyzer
Version: 2.0.0
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import docx
import openpyxl


logger = logging.getLogger(__name__)


class CRFSpecParser:
    """
    Parser for CRF (Case Report Form) specification documents.
    Supports DOCX and XLSX formats.
    """

    def __init__(self):
        """Initialize the parser."""
        self.parsed_data: Dict[str, Any] = {}
        self.variables: List[Dict[str, Any]] = []

    def parse(self, input_path: str) -> Dict[str, Any]:
        """
        Parse the CRF specification and extract structured data.

        Args:
            input_path: Path to the CRF specification (DOCX or XLSX).

        Returns:
            Dictionary containing parsed CRF information.
        """
        file_path = Path(input_path)
        file_ext = file_path.suffix.lower()

        if not file_path.exists():
            raise FileNotFoundError(f"CRF spec file not found: {file_path}")

        logger.info("Parsing CRF specification: %s", file_path)

        # Reset state for reentrant calls
        self.variables = []
        self.parsed_data = {}

        # Parse based on file type
        if file_ext == ".docx":
            self._parse_docx(file_path)
        elif file_ext in [".xlsx", ".xls"]:
            self._parse_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Build final data structure
        self.parsed_data = {
            "metadata": self._extract_metadata(file_path),
            "sections": self._organize_by_section(),
            "variables": self.variables,
            "variable_count": len(self.variables),
            "parsed_date": datetime.now().isoformat(),
        }

        logger.info(
            "Parsed CRF spec: variables=%d, sections=%d",
            self.parsed_data["variable_count"],
            len(self.parsed_data["sections"]),
        )

        return self.parsed_data

    # ------------------------------------------------------------------
    # Private parsing helpers
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: Path) -> None:
        """Parse DOCX CRF specification."""
        doc = docx.Document(file_path)

        current_section = "General"
        current_category = "General"

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Detect section headers (common patterns)
            if self._is_section_header(text):
                current_section = text
                current_category = text
            elif self._is_category_header(text):
                current_category = text

            # Try to extract variable information from the text
            var_info = self._parse_variable_from_text(text, current_section, current_category)
            if var_info:
                self.variables.append(var_info)

        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    var_info = self._parse_variable_from_table(cells, current_section)
                    if var_info:
                        self.variables.append(var_info)

    def _parse_xlsx(self, file_path: Path) -> None:
        """Parse XLSX CRF specification."""
        wb = openpyxl.load_workbook(file_path, data_only=True)

        # Try to find the variables sheet
        sheet = wb.active

        if sheet.max_row is None or sheet.max_row < 2:
            return

        # Get headers from first row
        headers: List[Any] = []
        for cell in sheet[1]:
            headers.append(cell.value)

        # Parse each row
        current_section = "General"

        for row_idx in range(2, sheet.max_row + 1):
            row_data: Dict[Any, Any] = {}
            for col_idx, header in enumerate(headers, start=1):
                cell = sheet.cell(row_idx, col_idx)
                row_data[header] = cell.value

            # Check for section change
            if "section" in row_data and row_data["section"]:
                current_section = str(row_data["section"])

            # Build variable info
            var_info = {
                "variable_name": (
                    row_data.get("variable_name")
                    or row_data.get("Variable")
                    or row_data.get("Field Name")
                    or f"var_{row_idx}"
                ),
                "label": (
                    row_data.get("label")
                    or row_data.get("Label")
                    or row_data.get("Description")
                    or ""
                ),
                "section": current_section,
                "data_type": row_data.get("data_type") or row_data.get("Type") or "text",
                "format": row_data.get("format") or row_data.get("Format") or "",
                "valid_range": row_data.get("valid_range") or row_data.get("Range") or "",
                "unit": row_data.get("unit") or row_data.get("Unit") or "",
                "required": row_data.get("required") or row_data.get("Required") or False,
                "notes": row_data.get("notes") or row_data.get("Notes") or "",
            }

            # Only add if has meaningful variable name
            if var_info["variable_name"] and var_info["variable_name"] != f"var_{row_idx}":
                self.variables.append(var_info)

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        section_patterns = [
            r"^Section\s*\d+",
            r"^[A-Z]\.\s+\w+",
            r"^[0-9]+\.\s+\w+",
            r"^(?:Demographics|Treatment|Medical\s*History|Laboratory|Response|Adverse\s*Events|Follow-up)",
        ]

        for pattern in section_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False

    def _is_category_header(self, text: str) -> bool:
        """Check if text is a category/subheader."""
        if len(text) < 50 and (text.isupper() or text.endswith(":")):
            return True
        return False

    def _parse_variable_from_text(
        self, text: str, section: str, category: str
    ) -> Optional[Dict[str, Any]]:
        """Parse variable information from plain text."""
        # Pattern: Variable Name (Label) - Type
        pattern = r"([A-Za-z_][A-Za-z0-9_]*)\s*[\(\[]?\s*(.+?)\s*[\)\]]?\s*[-:]\s*(\w+)"
        match = re.match(pattern, text)

        if match:
            return {
                "variable_name": match.group(1),
                "label": match.group(2).strip(),
                "section": section,
                "category": category,
                "data_type": match.group(3).strip().lower(),
                "format": "",
                "valid_range": "",
                "required": False,
                "notes": "",
            }
        return None

    def _parse_variable_from_table(
        self, cells: List[str], section: str
    ) -> Optional[Dict[str, Any]]:
        """Parse variable information from table row."""
        if len(cells) < 2:
            return None

        # Try to identify columns
        var_name = cells[0] if cells[0] else ""

        # Skip if doesn't look like a variable name
        if not var_name or not re.match(r"^[A-Za-z_]", var_name):
            return None

        return {
            "variable_name": var_name,
            "label": cells[1] if len(cells) > 1 else "",
            "section": section,
            "data_type": cells[2] if len(cells) > 2 else "text",
            "format": cells[3] if len(cells) > 3 else "",
            "valid_range": cells[4] if len(cells) > 4 else "",
            "required": False,
            "notes": cells[5] if len(cells) > 5 else "",
        }

    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract CRF metadata."""
        return {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "format": file_path.suffix.lower(),
            "parsed_date": datetime.now().isoformat(),
        }

    def _organize_by_section(self) -> Dict[str, List[Dict[str, Any]]]:
        """Organize variables by section."""
        sections: Dict[str, List[Dict[str, Any]]] = {}

        for var in self.variables:
            section = var.get("section", "General")
            if section not in sections:
                sections[section] = []
            sections[section].append(
                {
                    "variable_name": var.get("variable_name"),
                    "label": var.get("label"),
                    "data_type": var.get("data_type"),
                    "required": var.get("required"),
                }
            )

        return sections
