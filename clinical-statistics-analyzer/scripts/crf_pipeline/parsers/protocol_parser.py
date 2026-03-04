#!/usr/bin/env python3
"""
protocol_parser.py
Protocol Document Parser for Clinical Trial Protocols

Parses clinical trial protocol documents (DOCX/PDF) and extracts:
- Study metadata (ID, title, phase, sponsor)
- Disease/indication information
- Study design (randomization, blinding, arms)
- Endpoints (primary, secondary)
- Treatment arms
- Inclusion/exclusion criteria
- Sample size and statistical considerations

Author: Clinical Statistics Analyzer
Version: 2.0.0
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pdfplumber


logger = logging.getLogger(__name__)


class ProtocolParser:
    """
    Parser for clinical trial protocol documents.
    Supports DOCX and PDF formats.
    """

    def __init__(self, include_raw_text: bool = False):
        """
        Initialize the parser.

        Args:
            include_raw_text: If True, include truncated raw text in parse output.
        """
        self.include_raw_text = include_raw_text
        self.document_text: str = ""
        self.parsed_data: Dict[str, Any] = {}

    def parse(self, input_path: str) -> Dict[str, Any]:
        """
        Parse the protocol document and extract structured data.

        Args:
            input_path: Path to the protocol document (DOCX or PDF).

        Returns:
            Dictionary containing parsed protocol information.
        """
        file_path = Path(input_path)
        file_ext = file_path.suffix.lower()

        if not file_path.exists():
            raise FileNotFoundError(f"Protocol file not found: {file_path}")

        logger.info("Parsing protocol: %s", file_path)

        # Extract text based on file type
        if file_ext == ".docx":
            self.document_text = self._parse_docx(file_path)
        elif file_ext == ".pdf":
            self.document_text = self._parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Extract structured information
        self.parsed_data = {
            "metadata": self._extract_metadata(),
            "study_design": self._extract_study_design(),
            "disease_info": self._extract_disease_info(),
            "endpoints": self._extract_endpoints(),
            "treatment_arms": self._extract_treatment_arms(),
            "eligibility": self._extract_eligibility(),
            "statistics": self._extract_statistics(),
        }
        if self.include_raw_text:
            self.parsed_data["raw_text"] = self.document_text[:10000]

        logger.info(
            "Parsed protocol: study_id=%s, phase=%s, arms=%d",
            self.parsed_data["metadata"]["study_id"],
            self.parsed_data["metadata"]["phase"],
            len(self.parsed_data["treatment_arms"]),
        )

        return self.parsed_data

    # ------------------------------------------------------------------
    # Private parsing helpers
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)

        return "\n".join(text_parts)

    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        return "\n".join(text_parts)

    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract study metadata."""
        text = self.document_text

        # Study ID pattern
        study_id_pattern = r"(?:Protocol|Sponsor|Investigational)\s*(?:ID|Number|No\.?)\s*:?\s*([A-Z0-9\-]+)"
        study_id = self._search_pattern(study_id_pattern, text, group=1)

        # Title pattern
        title_pattern = r"(?:Title|Study\s*Title)[:\s]+(.+?)(?:\n|$)"
        title = self._search_pattern(title_pattern, text, group=1)

        # Phase pattern
        phase_pattern = r"(?:Phase)\s*(?:Study)?[:\s]*(I{1,3}|IV|1|2|3|4|a|b)"
        phase = self._search_pattern(phase_pattern, text, group=1)

        # Sponsor
        sponsor_pattern = r"(?:Sponsor)[:\s]+(.+?)(?:\n|$)"
        sponsor = self._search_pattern(sponsor_pattern, text, group=1)

        # Date patterns
        date_pattern = r"(?:Date|Ver(?:sion)?\.?\s*\d*)[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})"
        date = self._search_pattern(date_pattern, text, group=1)

        return {
            "study_id": study_id or "Unknown",
            "title": title or "Unknown",
            "phase": phase or "Unknown",
            "sponsor": sponsor or "Unknown",
            "version_date": date or "Unknown",
            "parsed_date": datetime.now().isoformat(),
        }

    def _extract_study_design(self) -> Dict[str, Any]:
        """Extract study design information."""
        text = self.document_text

        # Study type
        study_type = self._search_pattern(
            r"(?:Study\s*Design|Type\s*of\s*Study)[:\s]+(.+?)(?:\n|$)",
            text,
            group=1,
        )

        # Randomization
        randomization = self._search_pattern(
            r"(?:Randomization|Randomised)[:\s]*(.+?)(?:\n|$)",
            text,
            group=1,
        )

        # Blinding
        blinding = self._search_pattern(
            r"(?:Blinding|Blinded|Double[\-\s]?blind|Single[\-\s]?blind)[:\s]*(.+?)(?:\n|$)",
            text,
            group=1,
        )

        # Duration
        duration = self._search_pattern(
            r"(?:Duration|Treatment\s*Period|Study\s*Duration)[:\s]*(\d+\s*(?:weeks?|months?|years?|cycles?)?)",
            text,
            group=1,
        )

        return {
            "study_type": study_type or "Unknown",
            "randomization": randomization or "Not specified",
            "blinding": blinding or "Not specified",
            "duration": duration or "Not specified",
        }

    def _extract_disease_info(self) -> Dict[str, Any]:
        """Extract disease/indication information."""
        text = self.document_text

        # Disease/Indication
        disease = self._search_pattern(
            r"(?:Disease|Indication|Patient\s*Population)[:\s]+(.+?)(?:\n|$)",
            text,
            group=1,
        )

        # If not found in specific pattern, search for common disease patterns
        if not disease:
            disease = self._search_pattern(
                r"(?:Acute\s*Myeloid\s*Leukemia|AML|Acute\s*Lymphoblastic\s*Leukemia|ALL)",
                text,
            )

        return {
            "disease": disease or "Unknown",
            "disease_category": self._categorize_disease(disease),
        }

    def _categorize_disease(self, disease: Optional[str]) -> str:
        """Categorize the disease type."""
        if not disease:
            return "Unknown"

        disease_lower = disease.lower()

        if "aml" in disease_lower or "myeloid" in disease_lower:
            return "AML"
        elif "all" in disease_lower or "lymphoblastic" in disease_lower:
            return "ALL"
        elif "cml" in disease_lower or "myelogenous" in disease_lower:
            return "CML"
        elif "lymphoma" in disease_lower:
            return "Lymphoma"
        elif "myeloma" in disease_lower:
            return "Myeloma"
        else:
            return "Other"

    def _extract_endpoints(self) -> Dict[str, Any]:
        """Extract endpoint information."""
        text = self.document_text

        # Primary endpoint
        primary_patterns = [
            r"(?:Primary\s*Endpoint|Primary\s*Objective)[:\s]+(.+?)(?:\n\n|\n[A-Z])",
            r"(?:Primary)[:\s]+(.+?)(?:\n\n|\n[A-Z])",
        ]
        primary = self._search_patterns(primary_patterns, text, group=1)

        # Secondary endpoints
        secondary_patterns = [
            r"(?:Secondary\s*Endpoints|Secondary\s*Objectives)[:\s]+(.+?)(?:\n\n|\n[A-Z])",
            r"(?:Secondary)[:\s]+(.+?)(?:\n\n|\n[A-Z])",
        ]
        secondary = self._search_patterns(secondary_patterns, text, group=1)

        return {
            "primary": primary or "Unknown",
            "secondary": secondary or "Not specified",
        }

    def _extract_treatment_arms(self) -> List[Dict[str, Any]]:
        """Extract treatment arm information."""
        text = self.document_text
        arms = []

        # Look for common treatment arm structures
        arm_sections = re.finditer(
            r"(?:Treatment\s*Arms?|Study\s*Arms?|Cohorts?)[:\s]+(.+?)(?:\n\n|\n(?:Inclusion|Exclusion|Primary|Secondary|Statistical|$))",
            text,
            re.DOTALL | re.IGNORECASE,
        )

        for match in arm_sections:
            arm_text = match.group(1)

            # Extract individual arms
            individual_arms = re.split(
                r"(?:\n(?:\d+\.?)|(?<!\d)\n[A-Z]\.?|\n•|\n-)", arm_text
            )

            for arm in individual_arms:
                arm = arm.strip()
                if arm and len(arm) > 5:
                    arms.append(
                        {
                            "name": arm[:100],
                            "description": arm,
                            "type": self._classify_arm_type(arm),
                        }
                    )

        return arms if arms else [
            {"name": "Not specified", "description": "Not specified", "type": "Unknown"}
        ]

    def _classify_arm_type(self, arm_text: str) -> str:
        """Classify the treatment arm type."""
        text_lower = arm_text.lower()

        if "placebo" in text_lower:
            return "Placebo"
        elif "control" in text_lower:
            return "Control"
        elif "standard" in text_lower or "soc" in text_lower:
            return "Standard of Care"
        elif "low dose" in text_lower:
            return "Low Dose"
        elif "high dose" in text_lower:
            return "High Dose"
        elif "combination" in text_lower:
            return "Combination"
        elif "mono" in text_lower or "single" in text_lower:
            return "Monotherapy"
        else:
            return "Experimental"

    def _extract_eligibility(self) -> Dict[str, Any]:
        """Extract eligibility criteria."""
        text = self.document_text

        # Inclusion criteria
        inclusion_pattern = r"(?:Inclusion\s*Criteria|Inclusion\s*Eligibility)[:\s]+(.+?)(?:\n(?:Exclusion|$))"
        inclusion = self._search_pattern(inclusion_pattern, text, group=1)

        # Exclusion criteria
        exclusion_pattern = r"(?:Exclusion\s*Criteria|Exclusion\s*Eligibility)[:\s]+(.+?)(?:\n(?:Study|$))"
        exclusion = self._search_pattern(exclusion_pattern, text, group=1)

        return {
            "inclusion": inclusion or "Not specified",
            "exclusion": exclusion or "Not specified",
        }

    def _extract_statistics(self) -> Dict[str, Any]:
        """Extract statistical information."""
        text = self.document_text

        # Sample size
        sample_size = self._search_pattern(
            r"(?:Sample\s*Size|Number\s*of\s*Patients|Total\s*Patients)[:\s]*(\d+)",
            text,
            group=1,
        )

        # Power
        power = self._search_pattern(
            r"(?:Power|Statistical\s*Power)[:\s]*(\d+\%|\d+\.?\d*)",
            text,
            group=1,
        )

        # Alpha/Significance
        alpha = self._search_pattern(
            r"(?:Alpha|Significance\s*Level)[:\s]*(0\.\d+)",
            text,
            group=1,
        )

        return {
            "sample_size": int(sample_size) if sample_size else None,
            "power": power or "Not specified",
            "alpha": alpha or "Not specified",
        }

    def _search_pattern(
        self, pattern: str, text: str, group: int = 0
    ) -> Optional[str]:
        """Search for a pattern in text."""
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            if group > 0:
                return match.group(group).strip()
            return match.group(0).strip()
        return None

    def _search_patterns(
        self, patterns: List[str], text: str, group: int = 0
    ) -> Optional[str]:
        """Search for multiple patterns in text, return first match."""
        for pattern in patterns:
            result = self._search_pattern(pattern, text, group)
            if result:
                return result
        return None
