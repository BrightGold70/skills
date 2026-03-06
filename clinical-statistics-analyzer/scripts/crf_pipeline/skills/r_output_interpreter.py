"""
ROutputInterpreter — Tier 2 CSA Scientific Skill (KEY)

Reads R-generated CSV and DOCX outputs post-run.
Extracts StatValue-shaped dicts and writes data/*_stats.json sidecars.
Existing _write_hpw_manifest() reads those sidecars automatically —
zero modification to orchestrator required.

Maps to: exploratory-data-analysis OpenCode skill
"""

from __future__ import annotations

import datetime
import json
import logging
import re
from pathlib import Path
from typing import Optional

from ._base import CSASkillBase, CSASkillContext

logger = logging.getLogger(__name__)

# ── Patterns for DOCX text extraction ────────────────────────────────────────

# {script_basename: [(stat_key, regex_pattern, unit)]}
_DOCX_PATTERNS: dict[str, list[tuple[str, str, Optional[str]]]] = {
    "20_aml_eln_risk": [
        ("eln_favorable_pct",    r"[Ff]avorable[:\s]+(\d+\.?\d*)\s*%",        "percent"),
        ("eln_intermediate_pct", r"[Ii]ntermediate[:\s]+(\d+\.?\d*)\s*%",     "percent"),
        ("eln_adverse_pct",      r"[Aa]dverse[:\s]+(\d+\.?\d*)\s*%",          "percent"),
    ],
    "21_aml_composite_response": [
        ("cr_rate",  r"\bCR[:\s]+(\d+\.?\d*)\s*%",                            "percent"),
        ("cri_rate", r"\bCRi[:\s]+(\d+\.?\d*)\s*%",                           "percent"),
        ("ccr_rate", r"[Cc]omposite\s+(?:CR|cCR)[:\s]+(\d+\.?\d*)\s*%",      "percent"),
        ("orr",      r"ORR[:\s]+(\d+\.?\d*)\s*%",                             "percent"),
    ],
    "22_cml_tfr_analysis": [
        ("mmr_12mo", r"MMR\s+(?:at\s+)?12[:\s]+(\d+\.?\d*)\s*%",             "percent"),
        ("tfr_12mo", r"TFR\s+(?:at\s+)?12[:\s]+(\d+\.?\d*)\s*%",             "percent"),
        ("tfr_24mo", r"TFR\s+(?:at\s+)?24[:\s]+(\d+\.?\d*)\s*%",             "percent"),
    ],
    "23_cml_scores": [
        ("sokal_high_pct", r"Sokal\s+[Hh]igh[:\s]+(\d+\.?\d*)\s*%",          "percent"),
    ],
    "24_hct_gvhd_analysis": [
        ("agvhd_grade2_4_rate",        r"aGVHD\s+[Gg]rade\s+2[-–]4[:\s]+(\d+\.?\d*)\s*%",          "percent"),
        ("agvhd_grade3_4_rate",        r"aGVHD\s+[Gg]rade\s+3[-–]4[:\s]+(\d+\.?\d*)\s*%",          "percent"),
        ("cgvhd_moderate_severe_rate", r"(?:cGVHD|[Cc]hronic\s+GVHD)\s+[Mm]od[:\s]+(\d+\.?\d*)\s*%", "percent"),
        ("grfs_12mo",                  r"GRFS\s+(?:at\s+)?12[:\s]+(\d+\.?\d*)\s*%",                  "percent"),
    ],
    "25_aml_phase1_boin": [
        ("target_dlt_rate", r"[Tt]arget\s+DLT[:\s]+(\d+\.?\d*)\s*%",         "percent"),
    ],
    "05_safety": [
        ("ae_grade3plus_rate", r"[Gg]rade\s+[≥≧]?3\+?\s+AE[s]?[:\s]+(\d+\.?\d*)\s*%", "percent"),
    ],
    "02_table1": [
        ("n_total", r"[Nn]\s*=\s*(\d+)",                                       "patients"),
    ],
    "03_efficacy": [
        ("orr", r"ORR[:\s]+(\d+\.?\d*)\s*%",                                  "percent"),
    ],
}


def _stat_value(value: float, unit: Optional[str] = None,
                ci_lower: Optional[float] = None,
                ci_upper: Optional[float] = None,
                p_value: Optional[float] = None) -> dict:
    """Build a StatValue-compatible dict."""
    d: dict = {"value": value}
    if unit is not None:
        d["unit"] = unit
    if ci_lower is not None:
        d["ci_lower"] = ci_lower
    if ci_upper is not None:
        d["ci_upper"] = ci_upper
    if p_value is not None:
        d["p_value"] = p_value
    return d


class ROutputInterpreter(CSASkillBase):
    """
    Reads R-generated CSVs and DOCX outputs → writes data/*_stats.json.

    Priority order per script:
      1. CSV columns (structured, exact)
      2. DOCX regex fallback (when script produces only .docx)
    """

    def invoke(self, prompt: str, **kwargs) -> str:
        try:
            return f"[ROutputInterpreter] {prompt[:200]}"
        except Exception:
            return ""

    def interpret(self, output_dir: Path) -> "CSASkillContext":
        """
        Scan output_dir for R outputs. Extract key_statistics. Write sidecars.
        Updates self.context.key_statistics with all extracted values.
        Returns self.context.
        """
        output_dir = Path(output_dir)
        all_stats: dict = {}

        # ── CSV extraction ────────────────────────────────────────────────────
        all_stats.update(self._extract_cox_csv(output_dir))
        all_stats.update(self._extract_finegray_csv(output_dir))
        all_stats.update(self._extract_samplesize_csv(output_dir))

        # ── DOCX regex extraction ─────────────────────────────────────────────
        all_stats.update(self._extract_docx_all(output_dir))

        # Update context
        self.context.key_statistics.update(all_stats)

        # Write sidecar JSON consumed by _write_hpw_manifest()
        self._write_sidecar(output_dir, "r_output_interpreter", all_stats)

        self._log.info("ROutputInterpreter: extracted %d key_statistics", len(all_stats))
        return self.context

    # ── CSV extractors ────────────────────────────────────────────────────────

    def _extract_cox_csv(self, output_dir: Path) -> dict:
        """Parse Cox_*_Analysis.csv → os_hr, os_median_months."""
        stats: dict = {}
        try:
            for csv_path in sorted(output_dir.glob("Cox_*_Analysis.csv")):
                stats.update(self._parse_cox_file(csv_path))
        except Exception as exc:
            self._log.warning("Cox CSV extraction failed: %s", exc)
        return stats

    def _parse_cox_file(self, csv_path: Path) -> dict:
        stats: dict = {}
        try:
            import csv as csv_mod
            with open(csv_path, newline="", encoding="utf-8") as f:
                reader = csv_mod.DictReader(f)
                for row in reader:
                    var = row.get("variable", row.get("Variable", "")).upper()
                    if "OS" in var or "OVERALL" in var:
                        hr = self._to_float(row.get("hr", row.get("HR")))
                        if hr is not None:
                            stats["os_hr"] = _stat_value(
                                hr,
                                ci_lower=self._to_float(row.get("hr_lower", row.get("HR_lower"))),
                                ci_upper=self._to_float(row.get("hr_upper", row.get("HR_upper"))),
                                p_value=self._to_float(row.get("p_value", row.get("p.value"))),
                            )
                    if "MEDIAN" in var or "SURV" in var:
                        med = self._to_float(row.get("median", row.get("Median")))
                        if med is not None:
                            stats["os_median_months"] = _stat_value(
                                med, unit="months",
                                ci_lower=self._to_float(row.get("lower", row.get("ci_lower"))),
                                ci_upper=self._to_float(row.get("upper", row.get("ci_upper"))),
                            )
        except Exception as exc:
            self._log.warning("Failed to parse %s: %s", csv_path.name, exc)
        return stats

    def _extract_finegray_csv(self, output_dir: Path) -> dict:
        """Parse FineGray_*.csv → grfs_event_rate, agvhd_grade2_4_rate."""
        stats: dict = {}
        try:
            for csv_path in sorted(output_dir.glob("FineGray_*.csv")):
                stats.update(self._parse_finegray_file(csv_path))
        except Exception as exc:
            self._log.warning("FineGray CSV extraction failed: %s", exc)
        return stats

    def _parse_finegray_file(self, csv_path: Path) -> dict:
        stats: dict = {}
        try:
            import csv as csv_mod
            with open(csv_path, newline="", encoding="utf-8") as f:
                reader = csv_mod.DictReader(f)
                for row in reader:
                    cause = row.get("cause", row.get("Cause", "")).upper()
                    shr = self._to_float(row.get("shr", row.get("SHR")))
                    if shr is None:
                        continue
                    sv = _stat_value(
                        shr,
                        ci_lower=self._to_float(row.get("shr_lower")),
                        ci_upper=self._to_float(row.get("shr_upper")),
                        p_value=self._to_float(row.get("p_value", row.get("p.value"))),
                    )
                    if "GRFS" in cause:
                        stats["grfs_event_rate"] = sv
                    elif "AGVHD" in cause or "ACUTE" in cause:
                        stats["agvhd_grade2_4_rate"] = sv
                    elif "CGVHD" in cause or "CHRONIC" in cause:
                        stats["cgvhd_moderate_severe_rate"] = sv
        except Exception as exc:
            self._log.warning("Failed to parse %s: %s", csv_path.name, exc)
        return stats

    def _extract_samplesize_csv(self, output_dir: Path) -> dict:
        """Parse SampleSize_*.csv → n_total."""
        stats: dict = {}
        try:
            for csv_path in sorted(output_dir.glob("SampleSize_*.csv")):
                stats.update(self._parse_samplesize_file(csv_path))
                break  # first match only
        except Exception as exc:
            self._log.warning("SampleSize CSV extraction failed: %s", exc)
        return stats

    def _parse_samplesize_file(self, csv_path: Path) -> dict:
        stats: dict = {}
        try:
            import csv as csv_mod
            with open(csv_path, newline="", encoding="utf-8") as f:
                reader = csv_mod.DictReader(f)
                for row in reader:
                    param = row.get("parameter", row.get("Parameter", "")).lower()
                    if "n_total" in param or "n total" in param or "n_per_arm" in param:
                        val = self._to_float(row.get("value", row.get("Value")))
                        if val is not None:
                            stats["n_total"] = _stat_value(int(val), unit="patients")
        except Exception as exc:
            self._log.warning("Failed to parse %s: %s", csv_path.name, exc)
        return stats

    # ── DOCX regex extraction ─────────────────────────────────────────────────

    def _extract_docx_all(self, output_dir: Path) -> dict:
        """
        For each known DOCX output pattern, extract text and apply regex rules.
        Falls back gracefully if python-docx not installed.
        """
        stats: dict = {}
        try:
            import docx as _docx  # python-docx
            for script_key, patterns in _DOCX_PATTERNS.items():
                docx_files = list(output_dir.glob("*.docx"))
                if not docx_files:
                    break
                for docx_path in docx_files:
                    name_lower = docx_path.stem.lower()
                    # Match script key fragment to filename
                    key_fragment = script_key.split("_", 1)[-1] if "_" in script_key else script_key
                    if key_fragment.replace("_", " ") not in name_lower.replace("_", " "):
                        continue
                    text = self._docx_text(docx_path)
                    for stat_key, pattern, unit in patterns:
                        if stat_key in stats:
                            continue
                        m = re.search(pattern, text)
                        if m:
                            val = float(m.group(1))
                            # Convert percent string to fraction if > 1
                            if unit == "percent" and val > 1.0:
                                val = val / 100.0
                            stats[stat_key] = _stat_value(val, unit=unit)
        except ImportError:
            self._log.debug("python-docx not available — skipping DOCX extraction")
        except Exception as exc:
            self._log.warning("DOCX extraction failed: %s", exc)
        return stats

    def _docx_text(self, path: Path) -> str:
        try:
            import docx as _docx
            doc = _docx.Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return ""

    # ── Sidecar writer ────────────────────────────────────────────────────────

    def _write_sidecar(self, output_dir: Path, script_name: str, stats: dict) -> None:
        """
        Write data/{script_name}_stats.json in the format expected by
        _write_hpw_manifest(): {"key_statistics": ..., "disease_specific": {},
        "analysis_notes": {...}}.
        """
        try:
            data_dir = output_dir / "data"
            data_dir.mkdir(parents=True, exist_ok=True)
            sidecar = {
                "key_statistics": stats,
                "disease_specific": {},
                "analysis_notes": {
                    "source": script_name,
                    "extracted_at": datetime.datetime.utcnow().isoformat() + "Z",
                    "n_keys": len(stats),
                },
            }
            path = data_dir / f"{script_name}_stats.json"
            path.write_text(json.dumps(sidecar, indent=2, default=str))
            self._log.info("Wrote sidecar: %s (%d keys)", path.name, len(stats))
        except Exception as exc:
            self._log.warning("Failed to write sidecar: %s", exc)

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _to_float(val) -> Optional[float]:
        if val is None:
            return None
        try:
            return float(str(val).strip())
        except (ValueError, TypeError):
            return None
