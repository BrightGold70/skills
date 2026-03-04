"""DOCX processor for Word-format CRF documents."""

import logging
import os
import re
from pathlib import Path
from typing import List, Optional

from .base import DocumentProcessor, DocumentResult

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxProcessor(DocumentProcessor):
    """Process DOCX files for data extraction."""

    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in (".docx", ".doc")

    def process(self, file_path: str, hospital: str = "") -> DocumentResult:
        """Extract text from a single DOCX file."""
        result = DocumentResult(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            hospital=hospital,
            is_scanned=False,
            text="",
            format="docx",
        )

        if not DOCX_AVAILABLE:
            result.error = "python-docx not installed"
            return result

        try:
            doc = Document(file_path)
            parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)

            result.text = "\n".join(parts)
            logger.info("Extracted %d chars from %s", len(result.text), file_path)
        except Exception as e:
            logger.error("Error processing DOCX %s: %s", file_path, e)
            result.error = str(e)

        return result

    def process_directory(self, input_dir: str) -> List[DocumentResult]:
        """Walk subdirectories (hospital -> patient DOCX files)."""
        results = []
        input_path = Path(input_dir)

        for hospital_dir in sorted(input_path.iterdir()):
            if not hospital_dir.is_dir():
                continue
            hospital_name = hospital_dir.name
            logger.info("Processing hospital: %s", hospital_name)

            for docx_file in sorted(hospital_dir.glob("*.docx")):
                if "template" in docx_file.name.lower():
                    logger.info("  Skipping template: %s", docx_file.name)
                    continue
                logger.info("  Processing: %s", docx_file.name)
                results.append(self.process(str(docx_file), hospital=hospital_name))

        return results
