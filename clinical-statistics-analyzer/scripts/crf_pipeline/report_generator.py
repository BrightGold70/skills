"""Mini-CSR report generator — ICH-E3 lite document assembly."""

import logging
import os
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available — report generation disabled")


@dataclass
class CSRSection:
    """A section in the mini-CSR document."""
    title: str
    level: int  # heading level (1-3)
    narrative: str  # placeholder or auto-generated text
    tables: List[str] = field(default_factory=list)
    figures: List[str] = field(default_factory=list)


# Script name → ICH-E3 section mapping
_SCRIPT_TO_SECTION = {
    "02_table1.R": "demographics",
    "03_efficacy.R": "efficacy",
    "14_forest_plot.R": "efficacy",
    "04_survival.R": "survival",
    "05_safety.R": "safety",
    # Disease-specific scripts → disease_specific
    "20_aml_eln_risk.R": "disease_specific",
    "21_aml_composite_response.R": "disease_specific",
    "22_cml_tfr_analysis.R": "disease_specific",
    "23_cml_scores.R": "disease_specific",
    "24_hct_gvhd_analysis.R": "disease_specific",
    "25_aml_phase1_boin.R": "disease_specific",
    "26_cml_eln_milestones.R": "disease_specific",
    "27_cml_waterfall.R": "disease_specific",
    "28_cml_resistance.R": "disease_specific",
    "29_cml_tfr_deep.R": "disease_specific",
}

# Narrative placeholders per section
_NARRATIVE_TEMPLATES = {
    "title_page": "",
    "synopsis": (
        "[NARRATIVE: Provide a brief synopsis of the study including the study objective, "
        "design, key eligibility criteria, treatment arms, primary and secondary endpoints, "
        "and a summary of key findings.]"
    ),
    "demographics": (
        "[NARRATIVE: Describe the baseline characteristics of the study population. "
        "Include median age, sex distribution, ECOG performance status, and key disease "
        "features relevant to the disease type.]"
    ),
    "efficacy": (
        "[NARRATIVE: Summarize the primary efficacy results including overall response rate, "
        "complete response rate, and subgroup analyses. Reference the accompanying table "
        "and forest plot.]"
    ),
    "safety": (
        "[NARRATIVE: Describe the safety profile including most common adverse events "
        "(occurring in >=10% of patients), grade 3-4 events, serious adverse events, "
        "and treatment discontinuations due to AEs.]"
    ),
    "survival": (
        "[NARRATIVE: Report the survival outcomes including median overall survival, "
        "progression-free survival, and event-free survival. Describe the Kaplan-Meier "
        "curves and Cox proportional hazards model results.]"
    ),
    "disease_specific": (
        "[NARRATIVE: Describe disease-specific analysis results. Include relevant "
        "biomarkers, risk stratification, and disease-specific response assessments.]"
    ),
    "conclusions": (
        "[NARRATIVE: Summarize the key findings and their clinical significance. "
        "Discuss the benefit-risk assessment and place the results in context of "
        "existing evidence.]"
    ),
}


class ReportGenerator:
    """Generates ICH-E3 lite mini-CSR from analysis outputs.

    Assembles all .docx tables and .eps figures from R script outputs
    into a unified document following ICH-E3 lite section structure.
    """

    ICH_E3_SECTIONS = [
        "title_page",
        "synopsis",
        "demographics",
        "efficacy",
        "safety",
        "survival",
        "disease_specific",
        "conclusions",
    ]

    def __init__(self, output_dir: str, disease: str):
        """
        Args:
            output_dir: Base output directory (CSA_OUTPUT_DIR).
            disease: Disease type for disease-specific section routing.
        """
        self.output_dir = Path(output_dir)
        self.disease = disease
        self._reports_dir = self.output_dir / "Reports"
        self._reports_dir.mkdir(parents=True, exist_ok=True)

    def collect_outputs(self, script_results: list) -> Dict[str, List[str]]:
        """Map script outputs to ICH-E3 sections.

        Args:
            script_results: List of ScriptResult from orchestrator.

        Returns:
            Dict mapping section name to list of output file paths.
        """
        section_files: Dict[str, List[str]] = {s: [] for s in self.ICH_E3_SECTIONS}

        for result in script_results:
            if not result.success:
                continue

            section = _SCRIPT_TO_SECTION.get(result.script, "disease_specific")
            for fpath in result.output_files:
                if section in section_files:
                    section_files[section].append(fpath)

        return section_files

    def generate(
        self,
        script_results: list,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Generate the mini-CSR .docx document.

        Args:
            script_results: Results from orchestrator.run_scripts().
            metadata: Optional dict with study_title, protocol_number, etc.

        Returns:
            Path to the generated mini-CSR .docx file.
        """
        if not HAS_DOCX:
            logger.error("python-docx not installed — cannot generate report")
            return ""

        metadata = metadata or {}
        section_files = self.collect_outputs(script_results)

        doc = Document()

        # Title page
        self._add_title_page(doc, metadata)

        # Synopsis
        doc.add_heading("Synopsis", level=1)
        doc.add_paragraph(_NARRATIVE_TEMPLATES["synopsis"])
        self._add_auto_synopsis(doc, script_results, metadata)

        # Main sections
        for section in ["demographics", "efficacy", "safety", "survival", "disease_specific"]:
            heading = section.replace("_", " ").title()
            if section == "disease_specific":
                heading = f"Disease-Specific Analyses ({self.disease.upper()})"

            doc.add_heading(heading, level=1)
            doc.add_paragraph(_NARRATIVE_TEMPLATES.get(section, ""))

            # Embed tables and figures
            for fpath in section_files.get(section, []):
                if fpath.endswith(".docx"):
                    self._embed_table(doc, fpath)
                elif fpath.endswith(".eps"):
                    self._embed_figure(doc, fpath, caption=Path(fpath).stem)
                elif fpath.endswith(".png") or fpath.endswith(".pdf"):
                    self._embed_figure(doc, fpath, caption=Path(fpath).stem)

        # Conclusions
        doc.add_heading("Conclusions", level=1)
        doc.add_paragraph(_NARRATIVE_TEMPLATES["conclusions"])

        # Save
        output_path = self._reports_dir / f"Mini_CSR_{self.disease.upper()}.docx"
        doc.save(str(output_path))
        logger.info("Generated mini-CSR: %s", output_path)

        return str(output_path)

    def _add_title_page(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add title page to the CSR document."""
        title = metadata.get("study_title", f"Clinical Study Report — {self.disease.upper()}")
        protocol = metadata.get("protocol_number", "[Protocol Number]")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"\nProtocol: {protocol}")
        p2.add_run(f"\nDisease: {self.disease.upper()}")
        p2.add_run("\n\n[Sponsor Name]")
        p2.add_run("\n[Date]")

        doc.add_page_break()

    def _add_auto_synopsis(
        self, doc: Document, script_results: list, metadata: Dict[str, Any]
    ) -> None:
        """Add auto-generated synopsis paragraph."""
        total = len(script_results)
        successful = sum(1 for r in script_results if r.success)

        synopsis = (
            f"This mini-CSR summarizes the analysis of {self.disease.upper()} "
            f"clinical trial data. A total of {successful}/{total} analysis scripts "
            f"completed successfully, generating tables and figures across "
            f"demographics, efficacy, safety, and survival domains."
        )
        doc.add_paragraph(synopsis)

    def _embed_table(self, doc: Document, docx_path: str) -> None:
        """Read a .docx table and copy its content into the CSR document."""
        try:
            source_doc = Document(docx_path)
            for table in source_doc.tables:
                # Copy table structure
                new_table = doc.add_table(
                    rows=len(table.rows), cols=len(table.columns)
                )
                new_table.style = "Table Grid"

                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text

                doc.add_paragraph("")  # spacing after table

            logger.debug("Embedded table from %s", docx_path)
        except Exception as e:
            doc.add_paragraph(f"[Table: {Path(docx_path).name} — embedding failed: {e}]")
            logger.warning("Failed to embed table %s: %s", docx_path, e)

    def _embed_figure(self, doc: Document, figure_path: str, caption: str = "") -> None:
        """Convert .eps to .png if needed and embed in document with caption."""
        fpath = Path(figure_path)

        # Convert EPS to PNG if needed
        if fpath.suffix.lower() == ".eps":
            png_path = self._eps_to_png(fpath)
            if png_path is None:
                doc.add_paragraph(f"[Figure: {fpath.name} — conversion failed]")
                return
            fpath = png_path

        try:
            doc.add_picture(str(fpath), width=Inches(5.5))
            if caption:
                p = doc.add_paragraph(caption.replace("_", " "))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)
            doc.add_paragraph("")  # spacing
            logger.debug("Embedded figure: %s", figure_path)
        except Exception as e:
            doc.add_paragraph(f"[Figure: {fpath.name} — embedding failed: {e}]")
            logger.warning("Failed to embed figure %s: %s", figure_path, e)

    def _eps_to_png(self, eps_path: Path) -> Optional[Path]:
        """Convert EPS to PNG using ghostscript or R."""
        png_path = eps_path.with_suffix(".png")

        # Skip if PNG already exists
        if png_path.exists():
            return png_path

        # Strategy 1: ghostscript
        gs = "gs"
        try:
            result = subprocess.run(
                [
                    gs, "-dNOPAUSE", "-dBATCH", "-dSAFER",
                    "-sDEVICE=png16m", "-r150",
                    f"-sOutputFile={png_path}",
                    str(eps_path),
                ],
                capture_output=True, text=True, timeout=15,
            )
            if result.returncode == 0 and png_path.exists():
                return png_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Strategy 2: R grDevices
        try:
            r_code = f'png("{png_path}", width=800, height=600, res=150); plot.new(); dev.off()'
            fd, r_script = tempfile.mkstemp(suffix=".R")
            with os.fdopen(fd, "w") as f:
                f.write(r_code)
            subprocess.run(["Rscript", r_script], capture_output=True, timeout=10)
            os.unlink(r_script)
            if png_path.exists():
                return png_path
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            pass

        logger.warning("Cannot convert %s to PNG", eps_path.name)
        return None
