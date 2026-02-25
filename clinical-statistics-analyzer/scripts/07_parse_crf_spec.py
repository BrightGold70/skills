#!/usr/bin/env python3
"""
07_parse_crf_spec.py
CRF (Case Report Form) Specification Parser

Parses CRF specification documents (DOCX/XLSX) and extracts:
- Variable names and labels
- Data types and formats
- Valid ranges and value lists
- Skip patterns and dependencies
- Section/category information

Author: Clinical Statistics Analyzer
Version: 1.0.0
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Union

# Try to import required libraries
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

try:
    import openpyxl
except ImportError:
    print("Installing openpyxl...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl", "-q"])
    import openpyxl


class CRFSpecParser:
    """
    Parser for CRF (Case Report Form) specification documents.
    Supports DOCX and XLSX formats.
    """
    
    def __init__(self, file_path: str):
        """
        Initialize the parser with a CRF spec file.
        
        Args:
            file_path: Path to the CRF specification (DOCX or XLSX)
        """
        self.file_path = Path(file_path)
        self.file_ext = self.file_path.suffix.lower()
        self.parsed_data = {}
        self.variables = []
        
    def parse(self) -> Dict[str, Any]:
        """
        Parse the CRF specification and extract structured data.
        
        Returns:
            Dictionary containing parsed CRF information
        """
        if not self.file_path.exists():
            raise FileNotFoundError(f"CRF spec file not found: {self.file_path}")
        
        # Parse based on file type
        if self.file_ext == '.docx':
            self._parse_docx()
        elif self.file_ext in ['.xlsx', '.xls']:
            self._parse_xlsx()
        else:
            raise ValueError(f"Unsupported file format: {self.file_ext}")
        
        # Build final data structure
        self.parsed_data = {
            "metadata": self._extract_metadata(),
            "sections": self._organize_by_section(),
            "variables": self.variables,
            "variable_count": len(self.variables),
            "parsed_date": datetime.now().isoformat()
        }
        
        return self.parsed_data
    
    def _parse_docx(self) -> None:
        """Parse DOCX CRF specification."""
        doc = docx.Document(self.file_path)
        
        current_section = "General"
        current_category = "General"
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect section headers (common patterns)
            if self._is_section_header(text):
                current_section = text
                current_category = text
            elif self._is_category_header(text):
                current_category = text
            
            # Try to extract variable information from the text
            var_info = self._parse_variable_from_text(text, current_section, current_category)
            if var_info:
                self.variables.append(var_info)
        
        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    var_info = self._parse_variable_from_table(cells, current_section)
                    if var_info:
                        self.variables.append(var_info)
    
    def _parse_xlsx(self) -> None:
        """Parse XLSX CRF specification."""
        wb = openpyxl.load_workbook(self.file_path, data_only=True)
        
        # Try to find the variables sheet
        sheet = wb.active
        
        # Get headers from first row
        headers = []
        for cell in sheet[1]:
            headers.append(cell.value)
        
        # Parse each row
        current_section = "General"
        
        for row_idx in range(2, sheet.max_row + 1):
            row_data = {}
            for col_idx, header in enumerate(headers, start=1):
                cell = sheet.cell(row_idx, col_idx)
                row_data[header] = cell.value
            
            # Check for section change
            if 'section' in row_data and row_data['section']:
                current_section = str(row_data['section'])
            
            # Build variable info
            var_info = {
                "variable_name": row_data.get('variable_name') or row_data.get('Variable') or row_data.get('Field Name') or f"var_{row_idx}",
                "label": row_data.get('label') or row_data.get('Label') or row_data.get('Description') or "",
                "section": current_section,
                "data_type": row_data.get('data_type') or row_data.get('Type') or "text",
                "format": row_data.get('format') or row_data.get('Format') or "",
                "valid_range": row_data.get('valid_range') or row_data.get('Range') or "",
                "unit": row_data.get('unit') or row_data.get('Unit') or "",
                "required": row_data.get('required') or row_data.get('Required') or False,
                "notes": row_data.get('notes') or row_data.get('Notes') or ""
            }
            
            # Only add if has meaningful variable name
            if var_info["variable_name"] and var_info["variable_name"] != f"var_{row_idx}":
                self.variables.append(var_info)
    
    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        # Common section patterns
        section_patterns = [
            r'^Section\s*\d+',
            r'^[A-Z]\.\s+\w+',
            r'^[0-9]+\.\s+\w+',
            r'^(?:Demographics|Treatment|Medical\s*History|Laboratory|Response|Adverse\s*Events|Follow-up)',
        ]
        
        for pattern in section_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _is_category_header(self, text: str) -> bool:
        """Check if text is a category/subheader."""
        # Check for shorter headers that might be categories
        if len(text) < 50 and (text.isupper() or text.endswith(':')):
            return True
        return False
    
    def _parse_variable_from_text(self, text: str, section: str, category: str) -> Optional[Dict]:
        """Parse variable information from plain text."""
        # Pattern: Variable Name (Label) - Type
        pattern = r'([A-Za-z_][A-Za-z0-9_]*)\s*[\(\[]?\s*(.+?)\s*[\)\]]?\s*[-:]\s*(\w+)'
        match = re.match(pattern, text)
        
        if match:
            return {
                "variable_name": match.group(1),
                "label": match.group(2).strip(),
                "section": section,
                "category": category,
                "data_type": match.group(3).strip().lower(),
                "format": "",
                "valid_range": "",
                "required": False,
                "notes": ""
            }
        return None
    
    def _parse_variable_from_table(self, cells: List[str], section: str) -> Optional[Dict]:
        """Parse variable information from table row."""
        if len(cells) < 2:
            return None
        
        # Try to identify columns
        var_name = cells[0] if cells[0] else ""
        
        # Skip if doesn't look like a variable name
        if not var_name or not re.match(r'^[A-Za-z_]', var_name):
            return None
        
        return {
            "variable_name": var_name,
            "label": cells[1] if len(cells) > 1 else "",
            "section": section,
            "data_type": cells[2] if len(cells) > 2 else "text",
            "format": cells[3] if len(cells) > 3 else "",
            "valid_range": cells[4] if len(cells) > 4 else "",
            "required": False,
            "notes": cells[5] if len(cells) > 5 else ""
        }
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract CRF metadata."""
        return {
            "file_name": self.file_path.name,
            "file_path": str(self.file_path),
            "format": self.file_ext,
            "parsed_date": datetime.now().isoformat()
        }
    
    def _organize_by_section(self) -> Dict[str, List[Dict]]:
        """Organize variables by section."""
        sections = {}
        
        for var in self.variables:
            section = var.get('section', 'General')
            if section not in sections:
                sections[section] = []
            sections[section].append({
                "variable_name": var.get('variable_name'),
                "label": var.get('label'),
                "data_type": var.get('data_type'),
                "required": var.get('required')
            })
        
        return sections
    
    def save_json(self, output_path: Optional[str] = None) -> str:
        """
        Save parsed data to JSON file.
        
        Args:
            output_path: Path for output JSON file. If None, uses default naming.
            
        Returns:
            Path to saved JSON file
        """
        if not self.parsed_data:
            raise ValueError("No parsed data available. Run parse() first.")
        
        if output_path is None:
            output_path = self.file_path.parent / f"{self.file_path.stem}_spec_parsed.json"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.parsed_data, f, indent=2, ensure_ascii=False)
        
        return str(output_path)
    
    def save_csv(self, output_path: Optional[str] = None) -> str:
        """
        Save variables to CSV file.
        
        Args:
            output_path: Path for output CSV file.
            
        Returns:
            Path to saved CSV file
        """
        import csv
        
        if not self.variables:
            raise ValueError("No variables parsed. Run parse() first.")
        
        if output_path is None:
            output_path = self.file_path.parent / f"{self.file_path.stem}_variables.csv"
        
        fieldnames = ['variable_name', 'label', 'section', 'category', 'data_type', 
                      'format', 'valid_range', 'unit', 'required', 'notes']
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(self.variables)
        
        return str(output_path)


def main():
    """Main function for command-line usage."""
    if len(sys.argv) < 2:
        print("Usage: python 07_parse_crf_spec.py <crf_spec.docx/xlsx> [output_json] [output_csv]")
        print("\nExample:")
        print("  python 07_parse_crf_spec.py crf_spec.xlsx")
        print("  python 07_parse_crf_spec.py crf_spec.xlsx output.json variables.csv")
        sys.exit(1)
    
    crf_file = sys.argv[1]
    output_json = sys.argv[2] if len(sys.argv) > 2 else None
    output_csv = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        print(f"Parsing CRF specification: {crf_file}")
        parser = CRFSpecParser(crf_file)
        data = parser.parse()
        
        # Print summary
        print("\n" + "="*60)
        print("CRF Specification Parsed Successfully")
        print("="*60)
        print(f"Total Variables: {data['variable_count']}")
        print(f"Sections: {len(data['sections'])}")
        
        for section, vars_list in data['sections'].items():
            print(f"  - {section}: {len(vars_list)} variables")
        
        print("="*60)
        
        # Save to JSON
        json_path = parser.save_json(output_json)
        print(f"\nSaved JSON to: {json_path}")
        
        # Save to CSV
        if output_csv or len(sys.argv) <= 3:
            csv_path = parser.save_csv(output_csv)
            print(f"Saved CSV to: {csv_path}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
