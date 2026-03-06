"""
DocumentsTab: three-zone document management for CSA app.

Zone A — Protocol display (read from protocol_params.json via session state)
Zone B — CRF Form upload and schema preview
Zone C — Patient CRF batch upload, extraction, and validation
"""

import json
from pathlib import Path
from typing import Optional

import streamlit as st


class DocumentsTab:
    """Three-zone document upload and display panel."""

    def __init__(self, project_dir: str):
        self.project_dir = Path(project_dir)

    def render(self) -> None:
        self._render_zone_a()
        st.divider()
        self._render_zone_b()
        st.divider()
        self._render_zone_c()

    # ------------------------------------------------------------------
    # Zone A — Protocol (read-only from protocol_params.json)
    # ------------------------------------------------------------------

    def _render_zone_a(self) -> None:
        st.markdown("#### Zone A — Protocol Parameters")
        params = st.session_state.get("protocol_params")
        if not params:
            st.warning("No protocol_params.json found. Load a protocol in Phase 2 (Research Design) first.")
            return

        cols = st.columns(3)
        cols[0].metric("Study type", params.get("study_type") or "—")
        cols[1].metric("Sample size (N)", params.get("sample_size_n") or "—")
        cols[2].metric("Primary endpoint", params.get("primary_endpoint") or "—")

        if params.get("secondary_endpoints"):
            st.caption("Secondary endpoints: " + ", ".join(params["secondary_endpoints"]))
        if params.get("reporting_guideline"):
            st.caption(f"Reporting guideline: **{params['reporting_guideline']}**")
        if params.get("missing_fields"):
            st.warning("Missing fields: " + ", ".join(params["missing_fields"]))

    # ------------------------------------------------------------------
    # Zone B — CRF Form schema
    # ------------------------------------------------------------------

    def _render_zone_b(self) -> None:
        st.markdown("#### Zone B — CRF Form Schema")
        uploaded = st.file_uploader(
            "CRF Form (DOCX or XLSX)",
            type=["docx", "xlsx"],
            key="csa_crf_form_upload",
        )
        if uploaded:
            schema = self._parse_crf_form_schema(uploaded)
            if schema:
                st.session_state["crf_schema"] = schema
                st.success(f"Schema loaded: {len(schema)} fields")
                st.dataframe(
                    schema,
                    column_config={
                        "field": "Field",
                        "type": "Type",
                        "required": "Required",
                        "valid_values": "Valid values",
                    },
                    use_container_width=True,
                )
            else:
                st.error("Could not parse CRF form schema.")

        elif st.session_state.get("crf_schema"):
            st.caption(f"Schema in session: {len(st.session_state['crf_schema'])} fields")

    def _parse_crf_form_schema(self, file) -> list[dict]:
        """Parse CRF form DOCX or XLSX into field schema list."""
        name = file.name.lower()
        try:
            if name.endswith(".xlsx"):
                import pandas as pd
                df = pd.read_excel(file, sheet_name=0, header=0)
                df.columns = [c.strip().lower() for c in df.columns]
                col_map = {
                    "field": next((c for c in df.columns if "field" in c), df.columns[0]),
                    "type": next((c for c in df.columns if "type" in c), None),
                    "required": next((c for c in df.columns if "required" in c or "mandatory" in c), None),
                    "valid_values": next((c for c in df.columns if "value" in c or "option" in c), None),
                }
                rows = []
                for _, row in df.iterrows():
                    rows.append({
                        "field": str(row[col_map["field"]]).strip(),
                        "type": str(row[col_map["type"]]).strip() if col_map["type"] else "",
                        "required": str(row[col_map["required"]]).strip() if col_map["required"] else "",
                        "valid_values": str(row[col_map["valid_values"]]).strip() if col_map["valid_values"] else "",
                    })
                return [r for r in rows if r["field"] and r["field"] != "nan"]

            elif name.endswith(".docx"):
                from docx import Document
                doc = Document(file)
                if not doc.tables:
                    return []
                table = doc.tables[0]
                rows = []
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_dict = {
                        "field": cells[0] if len(cells) > 0 else "",
                        "type": cells[1] if len(cells) > 1 else "",
                        "required": cells[2] if len(cells) > 2 else "",
                        "valid_values": cells[3] if len(cells) > 3 else "",
                    }
                    if row_dict["field"]:
                        rows.append(row_dict)
                return rows
        except Exception as e:
            st.error(f"Schema parse error: {e}")
            return []
        return []

    # ------------------------------------------------------------------
    # Zone C — Patient CRFs (batch)
    # ------------------------------------------------------------------

    def _render_zone_c(self) -> None:
        st.markdown("#### Zone C — Patient CRFs (Batch)")
        schema = st.session_state.get("crf_schema")
        if not schema:
            st.info("Upload a CRF form schema in Zone B first, then upload patient CRFs here.")
            return

        uploaded_files = st.file_uploader(
            "Patient CRFs",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key="csa_patient_crfs",
        )

        if not uploaded_files:
            return

        if st.button("Process patient CRFs", key="process_crfs", type="primary"):
            status_rows = []
            progress = st.progress(0)
            for i, f in enumerate(uploaded_files):
                row = self._process_patient_crf(f, schema)
                status_rows.append(row)
                progress.progress((i + 1) / len(uploaded_files))

            progress.empty()
            st.dataframe(
                status_rows,
                column_config={
                    "filename": "File",
                    "status": "Status",
                    "n_fields_extracted": "Fields extracted",
                    "n_validation_errors": "Validation errors",
                },
                use_container_width=True,
            )

            success = [r for r in status_rows if r["status"] == "validated"]
            if success:
                self._write_consolidated_csv(status_rows)
                st.success(
                    f"{len(success)}/{len(status_rows)} files processed → "
                    f"`data/crf_consolidated.csv`"
                )

    def _process_patient_crf(self, file, schema: list[dict]) -> dict:
        """Extract and validate fields from a single patient CRF file."""
        result = {
            "filename": file.name,
            "status": "error",
            "n_fields_extracted": 0,
            "n_validation_errors": 0,
            "errors": [],
        }
        try:
            text = self._extract_text(file)
            extracted = self._extract_fields(text, schema)
            errors = self._validate_fields(extracted, schema)
            result.update({
                "status": "validated" if not errors else "validation_errors",
                "n_fields_extracted": len(extracted),
                "n_validation_errors": len(errors),
                "errors": errors,
            })
        except Exception as e:
            result["errors"] = [str(e)]
        return result

    def _extract_text(self, file) -> str:
        """Extract text from PDF or DOCX; OCR fallback for scanned PDFs."""
        name = file.name.lower()
        if name.endswith(".docx"):
            from docx import Document
            doc = Document(file)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append("  ".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(parts)

        elif name.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(file) as pdf:
                pages_text = [p.extract_text() or "" for p in pdf.pages]
            full_text = "\n".join(pages_text)
            if len(full_text) / max(len(pages_text), 1) < 100:
                try:
                    import pytesseract
                    from pdf2image import convert_from_bytes
                    images = convert_from_bytes(file.getvalue(), dpi=200)
                    full_text = "\n\n".join(pytesseract.image_to_string(img) for img in images)
                except ImportError:
                    pass
            return full_text
        return ""

    def _extract_fields(self, text: str, schema: list[dict]) -> dict:
        """Simple key:value extraction from CRF text."""
        import re
        extracted = {}
        for field_spec in schema:
            field_name = field_spec["field"]
            pattern = re.compile(
                rf"{re.escape(field_name)}\s*[:：]\s*([^\n\r]{{1,100}})",
                re.IGNORECASE,
            )
            m = pattern.search(text)
            if m:
                extracted[field_name] = m.group(1).strip()
        return extracted

    def _validate_fields(self, extracted: dict, schema: list[dict]) -> list[str]:
        """Check extracted values against schema valid_values."""
        errors = []
        for field_spec in schema:
            field_name = field_spec["field"]
            valid_raw = field_spec.get("valid_values", "")
            if not valid_raw or valid_raw in ("nan", ""):
                continue
            valid_values = [v.strip() for v in valid_raw.split(",")]
            value = extracted.get(field_name, "")
            if value and value not in valid_values:
                errors.append(f"{field_name}: '{value}' not in {valid_values}")
        return errors

    def _write_consolidated_csv(self, status_rows: list[dict]) -> None:
        """Write consolidated CSV stub to project data/."""
        out = self.project_dir / "data" / "crf_consolidated.csv"
        out.parent.mkdir(parents=True, exist_ok=True)
        import csv
        import io
        buf = io.StringIO()
        if status_rows:
            writer = csv.DictWriter(buf, fieldnames=["filename", "status", "n_fields_extracted", "n_validation_errors"])
            writer.writeheader()
            for row in status_rows:
                writer.writerow({k: row[k] for k in writer.fieldnames})
        out.write_text(buf.getvalue())
