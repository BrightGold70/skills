"""
Protocol Parser — extracts structured content from clinical study protocol documents.

Supports DOCX (python-docx) and PDF (pdfplumber + pytesseract OCR fallback).
Outputs:
  - introduction_seed.md      (Background + Objectives)
  - methods_seed.md           (Design + Eligibility + Treatment + Assessments)
  - statistical_methods_seed.md (SAP content)
  - protocol_params.json      (structured params for CSA auto-population)
  - protocol_extracted.json   (full extraction cache)
  - references.json           (PubMed-verified)
  - unverified_refs.txt       (failed PubMed lookup)
"""

from __future__ import annotations

import json
import re
import shutil
import time
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal, Optional


# ── Dataclasses ──────────────────────────────────────────────────────────────


@dataclass
class Section:
    heading: str
    heading_level: int  # 1 = chapter, 2 = section, 3 = subsection
    text: str
    page_start: int = 0  # PDF only
    category: str = ""  # assigned by SectionExtractor


@dataclass
class ProtocolReference:
    raw: str
    authors: str = ""
    title: str = ""
    journal: str = ""
    year: str = ""
    doi: str = ""
    pmid: str = ""
    verified: bool = False


@dataclass
class SAPParams:
    primary_endpoint: str = ""
    statistical_test: str = ""
    sample_size_n: Optional[int] = None
    power: Optional[float] = None
    alpha: Optional[float] = None
    dropout_rate: Optional[float] = None
    analysis_population: str = ""
    secondary_endpoints: list[str] = field(default_factory=list)
    missing_fields: list[str] = field(default_factory=list)


@dataclass
class ProtocolDocument:
    file_path: str
    file_type: Literal["docx", "pdf"]
    raw_text: str
    sections: list[Section] = field(default_factory=list)


@dataclass
class ExtractionResult:
    project_dir: str
    background_seed: str = ""
    methods_seed: str = ""
    statistical_methods_seed: str = ""
    sap_params: SAPParams = field(default_factory=SAPParams)
    references: list[ProtocolReference] = field(default_factory=list)
    reporting_guideline: str = ""
    study_type: str = ""
    disease_keywords: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# ── Section keyword taxonomy ──────────────────────────────────────────────────

SECTION_KEYWORDS: dict[str, list[str]] = {
    "background": [
        "background", "rationale", "introduction", "scientific background",
        "study rationale", "prior evidence", "unmet need", "disease background",
    ],
    "objectives": [
        "objective", "aim", "purpose", "hypothesis", "endpoint",
        "primary endpoint", "secondary endpoint", "study aim",
    ],
    "eligibility": [
        "eligibility", "inclusion criteria", "exclusion criteria",
        "inclusion/exclusion", "patient selection", "study population",
        "subject selection", "enrollment criteria",
    ],
    "study_design": [
        "study design", "design overview", "trial design", "overview",
        "randomization", "treatment arm", "dosing", "schedule",
        "dose escalation", "cohort",
    ],
    "assessments": [
        "assessment", "evaluation", "schedule of events",
        "response criteria", "efficacy assessment", "safety assessment",
        "tumor assessment", "disease assessment",
    ],
    "sap": [
        "statistical", "analysis plan", "sample size", "power",
        "statistical methods", "statistical analysis", "biostatistic",
    ],
    "references": [
        "reference", "bibliography", "literature cited", "citations",
    ],
}

HEMATOLOGY_DISEASE_KEYWORDS = [
    "AML", "CML", "MDS", "ALL", "CLL", "MPN", "lymphoma", "myeloma",
    "leukemia", "myeloid", "lymphoid", "GVHD", "HCT", "HSCT",
    "venetoclax", "azacitidine", "decitabine", "asciminib", "imatinib",
    "dasatinib", "ponatinib", "ruxolitinib", "gilteritinib", "midostaurin",
]


# ── DocxParser ────────────────────────────────────────────────────────────────


class DocxParser:
    """Parse DOCX protocol using python-docx."""

    def parse(self, file_path: str) -> ProtocolDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document(file_path)
        sections: list[Section] = []
        current_heading = "Preamble"
        current_level = 1
        current_paragraphs: list[str] = []

        # Iterate body elements in document order (paragraphs AND tables).
        # doc.paragraphs skips table cells — iterating doc.element.body children
        # preserves table content within the correct section context.
        from docx.text.paragraph import Paragraph as _DocxParagraph
        from docx.table import Table as _DocxTable

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = _DocxParagraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                level = self._detect_heading_level(para)
                if level:
                    if current_paragraphs:
                        sections.append(Section(
                            heading=current_heading,
                            heading_level=current_level,
                            text="\n\n".join(current_paragraphs),
                        ))
                    current_heading = text
                    current_level = level
                    current_paragraphs = []
                else:
                    current_paragraphs.append(text)

            elif tag == "tbl":
                # Extract all non-empty cell text from the table and
                # append to the current section's paragraph list.
                table = _DocxTable(child, doc)
                seen_cells: set[str] = set()
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # python-docx repeats merged cells; deduplicate within row
                        if cell_text and cell_text not in seen_cells:
                            seen_cells.add(cell_text)
                            row_texts.append(cell_text)
                    if row_texts:
                        current_paragraphs.append("  ".join(row_texts))

        # Final section
        if current_paragraphs:
            sections.append(Section(
                heading=current_heading,
                heading_level=current_level,
                text="\n\n".join(current_paragraphs),
            ))

        raw_text = "\n\n".join(
            f"{s.heading}\n{s.text}" for s in sections
        )
        return ProtocolDocument(
            file_path=file_path,
            file_type="docx",
            raw_text=raw_text,
            sections=sections,
        )

    def _detect_heading_level(self, para) -> Optional[int]:
        """Return heading level 1–3 or None if not a heading."""
        style_name = para.style.name if para.style else ""

        # 1. Named heading styles
        m = re.match(r"Heading (\d)", style_name, re.IGNORECASE)
        if m:
            return min(int(m.group(1)), 3)

        text = para.text.strip()
        if not text or len(text) > 120:
            return None

        # 2. Bold + short paragraph
        is_bold = all(run.bold for run in para.runs if run.text.strip())
        if is_bold and len(text) < 100:
            return 2

        # 3. Numbered section pattern: "1.", "1.1", "4.3.2"
        if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", text):
            dots = text.split(".")[0]
            return min(len(dots), 3)

        # 4. ALL CAPS short line
        if text.isupper() and len(text) < 80:
            return 1

        return None


# ── PdfParser ─────────────────────────────────────────────────────────────────


class PdfParser:
    """Parse PDF protocol using pdfplumber with pytesseract OCR fallback."""

    def parse(self, file_path: str) -> ProtocolDocument:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required: pip install pdfplumber")

        with pdfplumber.open(file_path) as pdf:
            pages = pdf.pages
            if self._is_scanned(pages):
                text, sections = self._ocr_extract(file_path)
            else:
                text, sections = self._text_extract(pages)

        return ProtocolDocument(
            file_path=file_path,
            file_type="pdf",
            raw_text=text,
            sections=sections,
        )

    def _is_scanned(self, pages) -> bool:
        sample = pages[:5]
        texts = [p.extract_text() or "" for p in sample]
        avg = sum(len(t) for t in texts) / max(len(sample), 1)
        return avg < 100

    def _text_extract(self, pages) -> tuple[str, list[Section]]:
        sections: list[Section] = []
        current_heading = "Preamble"
        current_level = 1
        current_lines: list[str] = []
        full_text_parts: list[str] = []

        for page_num, page in enumerate(pages):
            text = page.extract_text() or ""
            full_text_parts.append(text)
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                level = self._detect_pdf_heading(line, page)
                if level:
                    if current_lines:
                        sections.append(Section(
                            heading=current_heading,
                            heading_level=current_level,
                            text="\n".join(current_lines),
                            page_start=page_num + 1,
                        ))
                    current_heading = line
                    current_level = level
                    current_lines = []
                else:
                    current_lines.append(line)

        if current_lines:
            sections.append(Section(
                heading=current_heading,
                heading_level=current_level,
                text="\n".join(current_lines),
            ))

        return "\n\n".join(full_text_parts), sections

    def _detect_pdf_heading(self, line: str, page) -> Optional[int]:
        if len(line) > 120:
            return None
        if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", line):
            depth = line.split(".")[0].count(".") + 1
            return min(depth, 3)
        if line.isupper() and len(line) < 80:
            return 1
        return None

    def _ocr_extract(self, file_path: str) -> tuple[str, list[Section]]:
        try:
            import pytesseract
            from pdf2image import convert_from_path
        except ImportError:
            raise ImportError(
                "Scanned PDF detected. Install: pip install pytesseract pdf2image\n"
                "Also: brew install tesseract poppler"
            )

        images = convert_from_path(file_path, dpi=200)
        text_parts = []
        for img in images:
            text_parts.append(pytesseract.image_to_string(img))

        full_text = "\n\n".join(text_parts)
        # Simple section split for OCR output
        sections = [Section(
            heading="Full Text (OCR)",
            heading_level=1,
            text=full_text,
        )]
        return full_text, sections


# ── SectionExtractor ──────────────────────────────────────────────────────────


class SectionExtractor:
    """Assign categories to sections using keyword matching."""

    def categorize(self, sections: list[Section]) -> list[Section]:
        for section in sections:
            section.category = self._match_category(section.heading)
        # References are usually last — override if uncategorized tail sections
        self._fix_reference_tail(sections)
        return sections

    def _match_category(self, heading: str) -> str:
        h = heading.lower()
        for category, keywords in SECTION_KEYWORDS.items():
            if any(kw in h for kw in keywords):
                return category
        return "other"

    def _fix_reference_tail(self, sections: list[Section]) -> None:
        """Last uncategorized section containing numbered lines → references."""
        if not sections:
            return
        last = sections[-1]
        if last.category == "other":
            if re.search(r"^\s*\d+\.", last.text, re.MULTILINE):
                last.category = "references"


# ── BackgroundExtractor ───────────────────────────────────────────────────────


class BackgroundExtractor:
    def extract(self, sections: list[Section]) -> str:
        bg_sections = [s for s in sections if s.category == "background"]
        obj_sections = [s for s in sections if s.category == "objectives"]

        extracted_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        parts = [
            "<!-- AUTO-GENERATED FROM PROTOCOL — REVIEW BEFORE SUBMISSION -->",
            f"<!-- Extracted: {extracted_date} -->",
            "",
        ]

        if bg_sections:
            parts.append("## Background\n")
            parts.append("\n\n".join(s.text for s in bg_sections))
        else:
            parts.append("## Background\n\n[NOT FOUND IN PROTOCOL]")

        if obj_sections:
            parts.append("\n\n## Study Objectives\n")
            parts.append("\n\n".join(s.text for s in obj_sections))
        else:
            parts.append("\n\n## Study Objectives\n\n[NOT FOUND IN PROTOCOL]")

        return "\n".join(parts)


# ── MethodsExtractor ──────────────────────────────────────────────────────────


class MethodsExtractor:
    def extract(self, sections: list[Section]) -> tuple[str, str, str]:
        """
        Returns (methods_seed, study_type, reporting_guideline).
        """
        design_secs = [s for s in sections if s.category == "study_design"]
        elig_secs = [s for s in sections if s.category == "eligibility"]
        assess_secs = [s for s in sections if s.category == "assessments"]

        combined_text = " ".join(
            s.text.lower() for s in sections
        )

        # Try to extract verbatim study design from "Overall Study Design" sentence
        _design_text_raw = " ".join(s.text for s in design_secs) or " ".join(s.text for s in sections)
        _overall_match = re.search(
            r"(?:overall\s+study\s+design|study\s+design)\s*[:\;]\s*([^\n]{10,120})",
            _design_text_raw,
            re.IGNORECASE,
        )
        _verbatim_design = _overall_match.group(1).strip(" .") if _overall_match else ""

        # Detect study type + reporting guideline.
        # Phase I/II/III check MUST come before the generic "prospective/cohort" check
        # because a "prospective phase II study" is interventional, not observational.
        _phase_match = re.search(
            r"\bphase\s+(I{1,3}|[123]|IV|4)\b",
            combined_text,
            re.IGNORECASE,
        )
        if re.search(r"\brandomiz\w+\b.*\barms?\b|\barms?\b.*\brandomiz\w+\b", combined_text):
            study_type = _verbatim_design or "RCT"
            reporting_guideline = "CONSORT 2010"
        elif _phase_match:
            if _verbatim_design:
                study_type = _verbatim_design
            else:
                phase_num = _phase_match.group(1).upper()
                prefix = "Single-arm" if re.search(r"\bsingle[\s-]arm\b", combined_text, re.IGNORECASE) else "Interventional"
                study_type = f"{prefix} Phase {phase_num} Study"
            reporting_guideline = "CONSORT 2010"
        elif re.search(r"\bcohort\b|\bobservational\b|\bprospective\b|\bretrospective\b", combined_text):
            study_type = _verbatim_design or "Observational"
            reporting_guideline = "STROBE"
        elif re.search(r"\bcase\b.*\breport\b|\bcase report\b", combined_text):
            study_type = "Case Report"
            reporting_guideline = "CARE 2013"
        else:
            study_type = _verbatim_design or "Clinical Study"
            reporting_guideline = ""

        parts = [
            "<!-- AUTO-GENERATED FROM PROTOCOL -->",
            "",
            "## Study Design\n",
        ]

        if design_secs:
            parts.append("\n\n".join(s.text for s in design_secs))
        else:
            parts.append("[NOT FOUND IN PROTOCOL]")

        # Eligibility
        parts.append("\n\n## Eligibility Criteria\n")
        if elig_secs:
            for s in elig_secs:
                inclusion, exclusion = self._split_ie_criteria(s.text)
                if inclusion:
                    parts.append("### Inclusion Criteria\n")
                    parts.append(inclusion)
                if exclusion:
                    parts.append("\n### Exclusion Criteria\n")
                    parts.append(exclusion)
        else:
            parts.append("[NOT FOUND IN PROTOCOL]")

        # Treatment — extract paragraphs mentioning dose/arm/regimen from design sections
        treatment_lines = []
        treatment_keywords = {"dose", "dosing", "arm", "treatment", "regimen", "mg", "cycle", "infusion", "drug"}
        for s in design_secs:
            for para in s.text.split("\n\n"):
                if any(kw in para.lower() for kw in treatment_keywords):
                    treatment_lines.append(para.strip())
        parts.append("\n\n## Treatment\n")
        if treatment_lines:
            parts.append("\n\n".join(treatment_lines))
        else:
            parts.append("[NOT FOUND IN PROTOCOL]")

        # Assessments
        parts.append("\n\n## Assessments\n")
        if assess_secs:
            parts.append("\n\n".join(s.text for s in assess_secs))
        else:
            parts.append("[NOT FOUND IN PROTOCOL]")

        if reporting_guideline:
            parts.append(f"\n\n<!-- REPORTING GUIDELINE: {reporting_guideline} -->")

        return "\n".join(parts), study_type, reporting_guideline

    def _split_ie_criteria(self, text: str) -> tuple[str, str]:
        """Split combined I/E criteria text into inclusion and exclusion blocks."""
        inc_match = re.search(
            r"inclusion criteria[:\s]*(.*?)(?=exclusion criteria|$)",
            text, re.IGNORECASE | re.DOTALL
        )
        exc_match = re.search(
            r"exclusion criteria[:\s]*(.*?)$",
            text, re.IGNORECASE | re.DOTALL
        )
        inclusion = inc_match.group(1).strip() if inc_match else text
        exclusion = exc_match.group(1).strip() if exc_match else ""
        return inclusion, exclusion


# ── SAPExtractor ─────────────────────────────────────────────────────────────


class SAPExtractor:
    # Multiple sample-size patterns tried in priority order.
    # High-specificity (patient/subject context) first to avoid matching
    # "Phase 1", "10 mg", "10 sites", "a total of 10 visits", etc.
    _SAMPLE_SIZE_PATTERNS: list = [
        # "120 patients will be enrolled/randomized/recruited"
        re.compile(r"\b(\d{2,4})\s+(?:patients?|subjects?|participants?)\s+(?:will\s+be\s+)?(?:enrolled|randomized|recruited|included)", re.IGNORECASE),
        # "enroll/randomize 120 patients/subjects"
        re.compile(r"(?:enroll|randomize|recruit|include)\s+(?:a\s+total\s+of\s+)?(\d{2,4})\s+(?:patients?|subjects?|participants?)", re.IGNORECASE),
        # "planned sample size of 120" / "sample size of N=120"
        re.compile(r"(?:planned\s+)?sample\s+size\s+(?:of\s+)?(?:n\s*=\s*)?(\d{2,4})", re.IGNORECASE),
        # "a total of 120 patients" (requires patient noun after number)
        re.compile(r"(?:a\s+)?total\s+of\s+(\d{2,4})\s+(?:patients?|subjects?|participants?)", re.IGNORECASE),
        # "N = 120" standalone (last resort — most ambiguous)
        re.compile(r"\bn\s*=\s*(\d{2,4})\b", re.IGNORECASE),
    ]
    # Multiple power patterns tried in order (number-before and number-after)
    _POWER_PATTERNS: list = [
        re.compile(r"(\d{2,3})\s*%\s*power",                              re.IGNORECASE),  # "80% power"
        re.compile(r"power\s+of\s+(\d{2,3})\s*%?",                        re.IGNORECASE),  # "power of 80%"
        re.compile(r"power\s*[=:≥]\s*(\d{2,3})\s*%?",                    re.IGNORECASE),  # "power = 80%"
        re.compile(r"(\d{2,3})\s+percent\s+power",                        re.IGNORECASE),  # "80 percent power"
        re.compile(r"powered\s+(?:at|to\s+\w+\s+with)\s+(\d{2,3})\s*%?", re.IGNORECASE),  # "powered at 80%"
        re.compile(r"assuming\s+(\d{2,3})\s*%\s*power",                   re.IGNORECASE),  # "assuming 80% power"
        re.compile(r"power\s*(?:of\s*)?[=:]?\s*0\.(\d{2})",              re.IGNORECASE),  # "power of 0.80"
    ]
    ALPHA_PATTERN = re.compile(
        r"(?:alpha|α|significance level|type\s+I error)\s*(?:of\s*)?(?:=\s*)?([0-9.]+)",
        re.IGNORECASE
    )
    PRIMARY_ENDPOINT_PATTERN = re.compile(
        r"primary\s+(?:end[\s\-]?point|endpoint)"
        r"[:\s]+(?:is\s+|was\s+|the\s+|a\s+|will\s+be\s+)?"
        r"([A-Za-z][^\.\n]{4,120})",
        re.IGNORECASE,
    )
    SECONDARY_ENDPOINT_PATTERN = re.compile(
        r"secondary\s+(?:end\s*point|endpoint)[s]?[:\s]+([^\n]+)", re.IGNORECASE
    )
    DROPOUT_PATTERN = re.compile(
        r"(?:dropout|drop-out|loss to follow.up|attrition)\s*(?:rate\s*)?(?:of\s*)?(\d+)\s*%",
        re.IGNORECASE
    )
    TEST_KEYWORDS = [
        "log-rank", "log rank", "Cox", "Fine-Gray", "Fisher",
        "chi-square", "chi square", "Kaplan-Meier", "t-test", "ANOVA",
    ]
    POPULATION_KEYWORDS = {
        "intention-to-treat": ["intention-to-treat", "intent-to-treat", "ITT"],
        "per-protocol": ["per-protocol", "per protocol", "PP population"],
        "modified intention-to-treat": ["mITT", "modified ITT"],
        "safety": ["safety population", "treated patients"],
    }

    def extract(self, sections: list[Section]) -> tuple[str, SAPParams]:
        sap_secs = [s for s in sections if s.category == "sap"]
        sap_text = "\n\n".join(s.text for s in sap_secs) if sap_secs else ""

        # Also search all text for primary endpoint (may be in objectives)
        all_text = "\n\n".join(s.text for s in sections)

        params = SAPParams()

        # Sample size — try patterns in priority order (most specific first)
        _search_text = sap_text or all_text
        for _pat in self._SAMPLE_SIZE_PATTERNS:
            _m = _pat.search(_search_text)
            if _m:
                try:
                    params.sample_size_n = int(_m.group(1))
                    break
                except (ValueError, IndexError):
                    continue

        # Power — try explicit patterns first
        _search_text = sap_text or all_text
        for _pat in self._POWER_PATTERNS:
            _m = _pat.search(_search_text)
            if _m:
                try:
                    raw = _m.group(1)
                    if _pat == self._POWER_PATTERNS[-1]:
                        params.power = float(f"0.{raw}")
                    else:
                        val = int(raw)
                        params.power = val / 100.0 if val > 1 else float(val)
                    break
                except (ValueError, IndexError):
                    continue
        # Fallback: infer power from β (type II error): power = 1 − β
        if params.power is None:
            _beta_pat = re.compile(
                r"(?:β|beta|type\s+II\s+error)\s*[=:]\s*(0\.\d+|\.\d+|\d{1,2}\s*%)",
                re.IGNORECASE,
            )
            _bm = _beta_pat.search(_search_text)
            if _bm:
                try:
                    raw = _bm.group(1).strip().rstrip("%")
                    beta = float(raw) / 100.0 if float(raw) > 1 else float(raw)
                    params.power = round(1.0 - beta, 4)
                except ValueError:
                    pass

        # Alpha
        m = self.ALPHA_PATTERN.search(sap_text or all_text)
        if m:
            try:
                params.alpha = float(m.group(1))
            except ValueError:
                pass

        # Dropout
        m = self.DROPOUT_PATTERN.search(sap_text or all_text)
        if m:
            params.dropout_rate = int(m.group(1)) / 100.0

        # Primary endpoint — explicit label first
        m = self.PRIMARY_ENDPOINT_PATTERN.search(all_text)
        if m:
            params.primary_endpoint = m.group(1).strip().rstrip(".,;")
        # Fallback: infer from response-rate / efficacy context (Phase II single-arm designs)
        if not params.primary_endpoint:
            _ep_fallbacks = [
                re.compile(r"(overall\s+response\s+rate\b[^.\n]{0,60})",          re.IGNORECASE),
                re.compile(r"(objective\s+response\s+rate\b[^.\n]{0,60})",         re.IGNORECASE),
                re.compile(r"(ORR\b[^.\n]{0,40})",                                 re.IGNORECASE),
                re.compile(r"(complete\s+response\s+rate\b[^.\n]{0,60})",          re.IGNORECASE),
                re.compile(r"(overall\s+survival\b[^.\n]{0,60})",                  re.IGNORECASE),
                re.compile(r"(progression[- ]free\s+survival\b[^.\n]{0,60})",      re.IGNORECASE),
                re.compile(r"(?:final\s+)?response\s+rate\b([^.\n]{0,60})",        re.IGNORECASE),
            ]
            for _fp in _ep_fallbacks:
                _fm = _fp.search(all_text)
                if _fm:
                    params.primary_endpoint = _fm.group(1).strip().rstrip(".,;") or _fm.group(0).strip().rstrip(".,;")
                    break

        # Secondary endpoints
        for m in self.SECONDARY_ENDPOINT_PATTERN.finditer(all_text):
            ep = m.group(1).strip().rstrip(".,;")
            if ep and ep not in params.secondary_endpoints:
                params.secondary_endpoints.append(ep)

        # Statistical test
        for test in self.TEST_KEYWORDS:
            if test.lower() in (sap_text or all_text).lower():
                params.statistical_test = test
                break

        # Analysis population
        for pop_name, keywords in self.POPULATION_KEYWORDS.items():
            if any(kw.lower() in all_text.lower() for kw in keywords):
                params.analysis_population = pop_name
                break

        # Missing fields
        for field_name, value in [
            ("primary_endpoint", params.primary_endpoint),
            ("statistical_test", params.statistical_test),
            ("sample_size_n", params.sample_size_n),
            ("power", params.power),
            ("alpha", params.alpha),
        ]:
            if not value:
                params.missing_fields.append(field_name)

        # Build seed markdown
        seed = self._build_seed(sap_text, params)
        return seed, params

    def _build_seed(self, sap_text: str, params: SAPParams) -> str:
        parts = [
            "<!-- AUTO-GENERATED FROM PROTOCOL -->",
            "",
            "## Statistical Methods\n",
        ]

        if sap_text:
            parts.append(sap_text)
        else:
            parts.append("[SAP SECTION NOT FOUND IN PROTOCOL]")

        parts.append("\n\n### Extracted Parameters\n")
        parts.append(f"- **Primary endpoint**: {params.primary_endpoint or '[TO FILL]'}")
        parts.append(f"- **Statistical test**: {params.statistical_test or '[TO FILL]'}")
        parts.append(f"- **Sample size**: N = {params.sample_size_n or '[TO FILL]'}")
        parts.append(f"- **Power**: {int(params.power * 100) if params.power else '[TO FILL]'}%")
        parts.append(f"- **Alpha**: {params.alpha or '[TO FILL]'}")
        parts.append(f"- **Dropout rate**: {int(params.dropout_rate * 100) if params.dropout_rate else '[TO FILL]'}%")
        parts.append(f"- **Analysis population**: {params.analysis_population or '[TO FILL]'}")

        if params.secondary_endpoints:
            parts.append(f"- **Secondary endpoints**: {'; '.join(params.secondary_endpoints[:5])}")

        if params.missing_fields:
            parts.append(f"\n<!-- MISSING FIELDS: {', '.join(params.missing_fields)} -->")

        return "\n".join(parts)


# ── ReferenceImporter ─────────────────────────────────────────────────────────


class ReferenceImporter:
    VANCOUVER_PATTERN = re.compile(r"^\s*(\d+)\.\s+(.+)", re.MULTILINE)
    DOI_PATTERN = re.compile(r"doi:\s*(10\.\d{4,}/\S+)", re.IGNORECASE)
    YEAR_PATTERN = re.compile(r"\b(19|20)\d{2}\b")

    def extract_references(self, sections: list[Section]) -> list[ProtocolReference]:
        ref_secs = [s for s in sections if s.category == "references"]
        if not ref_secs:
            return []

        ref_text = "\n\n".join(s.text for s in ref_secs)
        refs = []

        # Try Vancouver numbered format first
        matches = list(self.VANCOUVER_PATTERN.finditer(ref_text))
        if matches:
            for m in matches:
                raw = m.group(2).strip()
                refs.append(self._parse_ref(raw))
        else:
            # Fallback: split by double newline or numbered lines
            for chunk in re.split(r"\n{2,}", ref_text):
                chunk = chunk.strip()
                if len(chunk) > 20:
                    refs.append(self._parse_ref(chunk))

        return refs

    def _parse_ref(self, raw: str) -> ProtocolReference:
        ref = ProtocolReference(raw=raw)

        # DOI
        m = self.DOI_PATTERN.search(raw)
        if m:
            ref.doi = m.group(1).rstrip(".")

        # Year
        m = self.YEAR_PATTERN.search(raw)
        if m:
            ref.year = m.group(0)

        # Title: heuristic — text between first period and journal name
        # (Vancouver: Authors. Title. Journal. Year;vol:pages)
        parts = raw.split(". ")
        if len(parts) >= 3:
            ref.authors = parts[0]
            ref.title = parts[1]
            ref.journal = parts[2].split(";")[0].split(",")[0]

        return ref

    def verify_with_pubmed(
        self,
        refs: list[ProtocolReference],
        verbose: bool = True,
    ) -> list[ProtocolReference]:
        """Attempt PubMed verification for each reference."""
        try:
            from .pubmed_verifier import PubMedVerifier
            verifier = PubMedVerifier()
        except Exception:
            # PubMed verifier unavailable — mark all as unverified
            return refs

        for i, ref in enumerate(refs):
            try:
                if verbose:
                    print(f"  Verifying ref {i+1}/{len(refs)}...", end="\r")
                result = verifier.search_by_title(ref.title) if ref.title else None
                if result:
                    ref.pmid = str(result.get("pmid", ""))
                    ref.verified = bool(ref.pmid)
                time.sleep(0.35)  # NCBI rate limit: max 3/s
            except Exception:
                pass

        if verbose:
            print()
        return refs


# ── ProtocolParser (public API) ───────────────────────────────────────────────


class ProtocolParser:
    """Main entry point for protocol extraction."""

    def __init__(self, project_dir: str):
        self.project_dir = Path(project_dir)
        self._ensure_dirs()

    def _ensure_dirs(self) -> None:
        for subdir in ["docs/protocol", "docs/drafts", "data", "literature"]:
            (self.project_dir / subdir).mkdir(parents=True, exist_ok=True)

    def load(self, file_path: str) -> ProtocolDocument:
        """Load DOCX or PDF and return ProtocolDocument."""
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".docx":
            doc = DocxParser().parse(file_path)
        elif suffix == ".pdf":
            doc = PdfParser().parse(file_path)
        else:
            raise ValueError(f"Unsupported format: {suffix}. Use .docx or .pdf")

        # Copy to project protocol dir
        dest = self.project_dir / "docs" / "protocol" / path.name
        if str(path.resolve()) != str(dest.resolve()):
            shutil.copy2(file_path, dest)

        return doc

    def extract_all(
        self,
        doc: ProtocolDocument,
        import_refs: bool = True,
        verify_refs: bool = True,
        verbose: bool = True,
    ) -> ExtractionResult:
        """Run all extractors. Write output files. Return ExtractionResult."""
        result = ExtractionResult(project_dir=str(self.project_dir))

        if verbose:
            print(f"  Parsed {len(doc.sections)} sections ({doc.file_type.upper()})")

        # Categorize sections
        sections = SectionExtractor().categorize(doc.sections)

        # Background
        result.background_seed = BackgroundExtractor().extract(sections)
        self._write(result.background_seed, "docs/drafts/introduction_seed.md")
        if verbose:
            print("  ✓ Background extracted → docs/drafts/introduction_seed.md")

        # Methods
        methods_seed, study_type, reporting_guideline = MethodsExtractor().extract(sections)
        result.methods_seed = methods_seed
        result.study_type = study_type
        result.reporting_guideline = reporting_guideline
        self._write(result.methods_seed, "docs/drafts/methods_seed.md")
        if verbose:
            print("  ✓ Methods extracted   → docs/drafts/methods_seed.md")

        # SAP
        sap_seed, sap_params = SAPExtractor().extract(sections)
        result.statistical_methods_seed = sap_seed
        result.sap_params = sap_params
        self._write(result.statistical_methods_seed, "docs/drafts/statistical_methods_seed.md")
        if verbose:
            print("  ✓ SAP extracted       → docs/drafts/statistical_methods_seed.md")

        # Disease keywords
        all_text = doc.raw_text
        result.disease_keywords = [
            kw for kw in HEMATOLOGY_DISEASE_KEYWORDS if kw.lower() in all_text.lower()
        ]

        # References
        if import_refs:
            importer = ReferenceImporter()
            refs = importer.extract_references(sections)
            if refs and verify_refs:
                if verbose:
                    print(f"  ✓ {len(refs)} references found → verifying with PubMed...")
                refs = importer.verify_with_pubmed(refs, verbose=verbose)
            result.references = refs

        # protocol_params.json
        self._write_params(result)
        if verbose:
            print("  ✓ Parameters saved    → data/protocol_params.json")

        # Full extraction cache
        self._write_extraction_cache(doc, result)

        # Literature outputs
        if import_refs and result.references:
            verified = [r for r in result.references if r.verified]
            unverified = [r for r in result.references if not r.verified]
            self._write_refs_json(verified)
            self._write_unverified(unverified)
            if verbose:
                print(
                    f"  ✓ {len(verified)} verified (PMID assigned)\n"
                    f"  {'  ' if not unverified else '⚠ '}"
                    f"{len(unverified)} unverified → literature/unverified_refs.txt"
                )

        if result.sap_params.missing_fields:
            w = f"Missing SAP fields: {', '.join(result.sap_params.missing_fields)}"
            result.warnings.append(w)
            if verbose:
                print(f"  ⚠ {w}")

        return result

    def load_and_extract(
        self,
        file_path: str,
        import_refs: bool = True,
        verify_refs: bool = True,
        verbose: bool = True,
    ) -> ExtractionResult:
        """Convenience: load + extract_all."""
        doc = self.load(file_path)
        return self.extract_all(doc, import_refs=import_refs, verify_refs=verify_refs, verbose=verbose)

    # ── Private helpers ────────────────────────────────────────────────────────

    def _write(self, content: str, rel_path: str) -> None:
        path = self.project_dir / rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")

    def _write_params(self, result: ExtractionResult) -> None:
        params = asdict(result.sap_params)
        protocol_params = {
            "version": "1.0",
            "extracted_at": datetime.now(timezone.utc).isoformat(),
            "source_file": Path(result.project_dir).name,
            "primary_endpoint": result.sap_params.primary_endpoint,
            "statistical_test": result.sap_params.statistical_test,
            "sample_size_n": result.sap_params.sample_size_n,
            "power": result.sap_params.power,
            "alpha": result.sap_params.alpha,
            "dropout_rate": result.sap_params.dropout_rate,
            "analysis_population": result.sap_params.analysis_population,
            "secondary_endpoints": result.sap_params.secondary_endpoints,
            "study_type": result.study_type,
            "reporting_guideline": result.reporting_guideline,
            "disease_keywords": result.disease_keywords,
            "missing_fields": result.sap_params.missing_fields,
        }
        path = self.project_dir / "data" / "protocol_params.json"
        path.write_text(json.dumps(protocol_params, indent=2), encoding="utf-8")

    def _write_extraction_cache(self, doc: ProtocolDocument, result: ExtractionResult) -> None:
        cache = {
            "extracted_at": datetime.now(timezone.utc).isoformat(),
            "file_path": doc.file_path,
            "file_type": doc.file_type,
            "section_count": len(doc.sections),
            "sections": [
                {"heading": s.heading, "category": s.category, "chars": len(s.text)}
                for s in doc.sections
            ],
            "study_type": result.study_type,
            "reporting_guideline": result.reporting_guideline,
            "disease_keywords": result.disease_keywords,
            "sap_params": asdict(result.sap_params),
            "reference_count": len(result.references),
            "verified_count": sum(1 for r in result.references if r.verified),
            "warnings": result.warnings,
        }
        path = self.project_dir / "data" / "protocol_extracted.json"
        path.write_text(json.dumps(cache, indent=2), encoding="utf-8")

    def _write_refs_json(self, refs: list[ProtocolReference]) -> None:
        data = [asdict(r) for r in refs]
        path = self.project_dir / "literature" / "references.json"
        path.write_text(json.dumps(data, indent=2), encoding="utf-8")

    def _write_unverified(self, refs: list[ProtocolReference]) -> None:
        if not refs:
            return
        lines = [r.raw for r in refs]
        path = self.project_dir / "literature" / "unverified_refs.txt"
        path.write_text("\n\n".join(lines), encoding="utf-8")
